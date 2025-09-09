"""
Book Builder Module

Manages outline management, chapter orchestration, and export to various formats.
Handles the complete book generation workflow.

Chosen libraries:
- asyncio: Asynchronous book building operations
- pydantic: Data validation and type safety
- logging: Book building activity logging
- markdown: Markdown export functionality
- docx: DOCX export functionality

Adapted from: LangGraph (https://github.com/langchain-ai/langgraph)
Pattern: Stateful workflow orchestration with persistence
"""

import asyncio
import json
import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

import pydantic

logger = logging.getLogger(__name__)


class BookOutline(pydantic.BaseModel):
    """Model for book outlines."""
    book_id: str
    title: str
    subtitle: Optional[str] = None
    author: str
    description: str
    target_audience: str
    estimated_word_count: int
    chapters: List[Dict[str, Any]] = []
    created_at: datetime
    updated_at: datetime


class BookChapter(pydantic.BaseModel):
    """Model for book chapters."""
    chapter_id: str
    book_id: str
    title: str
    order: int
    content: str
    word_count: int
    status: str = "draft"  # draft, in_progress, completed, reviewed
    created_at: datetime
    updated_at: datetime
    metadata: Dict[str, Any] = {}


class BookManuscript(pydantic.BaseModel):
    """Model for complete book manuscripts."""
    book_id: str
    title: str
    subtitle: Optional[str] = None
    author: str
    description: str
    target_audience: str
    total_word_count: int
    chapters: List[BookChapter] = []
    created_at: datetime
    updated_at: datetime
    status: str = "draft"  # draft, in_progress, completed, published
    metadata: Dict[str, Any] = {}


class BookBuilder:
    """
    Manages book building, formatting, and export operations.
    
    Responsibilities:
    - Book outline management
    - Chapter orchestration and tracking
    - Export to multiple formats (Markdown, DOCX, PDF)
    - Progress tracking and persistence
    - Quality assurance and validation
    """
    
    def __init__(self, output_directory: str = "./books"):
        """
        Initialize book builder.
        
        Args:
            output_directory: Directory for book outputs
        """
        self.output_directory = Path(output_directory)
        self.output_directory.mkdir(parents=True, exist_ok=True)
        
        # Book state tracking
        self.current_books: Dict[str, BookManuscript] = {}
        self.book_outlines: Dict[str, BookOutline] = {}
        
        logger.info(f"Book builder initialized with output directory: {self.output_directory}")
    
    async def create_book_outline(
        self,
        title: str,
        author: str,
        description: str,
        target_audience: str,
        estimated_word_count: int = 50000,
        subtitle: Optional[str] = None
    ) -> BookOutline:
        """
        Create a new book outline.
        
        Args:
            title: Book title
            author: Book author
            description: Book description
            target_audience: Target audience
            estimated_word_count: Estimated total word count
            subtitle: Optional subtitle
            
        Returns:
            Created book outline
        """
        book_id = self._generate_book_id(title)
        
        outline = BookOutline(
            book_id=book_id,
            title=title,
            subtitle=subtitle,
            author=author,
            description=description,
            target_audience=target_audience,
            estimated_word_count=estimated_word_count,
            created_at=datetime.now(),
            updated_at=datetime.now()
        )
        
        self.book_outlines[book_id] = outline
        
        # Create book directory
        book_dir = self.output_directory / book_id
        book_dir.mkdir(exist_ok=True)
        
        # Save outline
        await self._save_outline(outline)
        
        logger.info(f"Created book outline: {title} (ID: {book_id})")
        
        return outline
    
    async def add_chapter_to_outline(
        self,
        book_id: str,
        chapter_title: str,
        chapter_description: str,
        estimated_word_count: int = 3000,
        order: Optional[int] = None
    ) -> bool:
        """
        Add a chapter to book outline.
        
        Args:
            book_id: Book ID
            chapter_title: Chapter title
            chapter_description: Chapter description
            estimated_word_count: Estimated word count for chapter
            order: Chapter order (auto-assigned if None)
            
        Returns:
            True if successful
        """
        if book_id not in self.book_outlines:
            logger.error(f"Book outline not found: {book_id}")
            return False
        
        outline = self.book_outlines[book_id]
        
        if order is None:
            order = len(outline.chapters) + 1
        
        chapter_data = {
            "chapter_id": f"{book_id}_ch_{order:02d}",
            "title": chapter_title,
            "description": chapter_description,
            "estimated_word_count": estimated_word_count,
            "order": order,
            "status": "planned"
        }
        
        outline.chapters.append(chapter_data)
        outline.updated_at = datetime.now()
        
        # Save updated outline
        await self._save_outline(outline)
        
        logger.info(f"Added chapter to outline: {chapter_title}")
        
        return True
    
    async def start_book_manuscript(self, book_id: str) -> BookManuscript:
        """
        Start a book manuscript from outline.
        
        Args:
            book_id: Book ID
            
        Returns:
            Created book manuscript
        """
        if book_id not in self.book_outlines:
            raise ValueError(f"Book outline not found: {book_id}")
        
        outline = self.book_outlines[book_id]
        
        manuscript = BookManuscript(
            book_id=book_id,
            title=outline.title,
            subtitle=outline.subtitle,
            author=outline.author,
            description=outline.description,
            target_audience=outline.target_audience,
            total_word_count=0,
            created_at=datetime.now(),
            updated_at=datetime.now()
        )
        
        self.current_books[book_id] = manuscript
        
        # Save manuscript
        await self._save_manuscript(manuscript)
        
        logger.info(f"Started book manuscript: {outline.title}")
        
        return manuscript
    
    async def add_chapter_content(
        self,
        book_id: str,
        chapter_id: str,
        content: str,
        status: str = "completed"
    ) -> bool:
        """
        Add chapter content to manuscript.
        
        Args:
            book_id: Book ID
            chapter_id: Chapter ID
            content: Chapter content
            status: Chapter status
            
        Returns:
            True if successful
        """
        if book_id not in self.current_books:
            logger.error(f"Book manuscript not found: {book_id}")
            return False
        
        manuscript = self.current_books[book_id]
        
        # Find chapter in outline
        chapter_data = None
        for ch in manuscript.chapters:
            if ch.chapter_id == chapter_id:
                chapter_data = ch
                break
        
        if not chapter_data:
            # Create new chapter
            chapter_data = BookChapter(
                chapter_id=chapter_id,
                book_id=book_id,
                title=f"Chapter {len(manuscript.chapters) + 1}",
                order=len(manuscript.chapters) + 1,
                content=content,
                word_count=len(content.split()),
                status=status,
                created_at=datetime.now(),
                updated_at=datetime.now()
            )
            manuscript.chapters.append(chapter_data)
        else:
            # Update existing chapter
            chapter_data.content = content
            chapter_data.word_count = len(content.split())
            chapter_data.status = status
            chapter_data.updated_at = datetime.now()
        
        # Update manuscript
        manuscript.total_word_count = sum(ch.word_count for ch in manuscript.chapters)
        manuscript.updated_at = datetime.now()
        
        # Save manuscript
        await self._save_manuscript(manuscript)
        
        logger.info(f"Added chapter content: {chapter_id} ({chapter_data.word_count} words)")
        
        return True
    
    async def export_to_markdown(
        self,
        book_id: str,
        output_path: Optional[str] = None
    ) -> str:
        """
        Export book to Markdown format.
        
        Args:
            book_id: Book ID
            output_path: Optional output path
            
        Returns:
            Path to exported file
        """
        if book_id not in self.current_books:
            raise ValueError(f"Book manuscript not found: {book_id}")
        
        manuscript = self.current_books[book_id]
        
        if not output_path:
            output_path = self.output_directory / book_id / f"{book_id}.md"
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Generate Markdown content
        markdown_content = self._generate_markdown_content(manuscript)
        
        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        
        logger.info(f"Exported book to Markdown: {output_path}")
        
        return str(output_path)
    
    async def export_to_docx(
        self,
        book_id: str,
        output_path: Optional[str] = None
    ) -> str:
        """
        Export book to DOCX format.
        
        Args:
            book_id: Book ID
            output_path: Optional output path
            
        Returns:
            Path to exported file
        """
        if book_id not in self.current_books:
            raise ValueError(f"Book manuscript not found: {book_id}")
        
        manuscript = self.current_books[book_id]
        
        if not output_path:
            output_path = self.output_directory / book_id / f"{book_id}.docx"
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Generate DOCX content
        await self._generate_docx_content(manuscript, output_path)
        
        logger.info(f"Exported book to DOCX: {output_path}")
        
        return str(output_path)
    
    async def export_to_pdf(
        self,
        book_id: str,
        output_path: Optional[str] = None
    ) -> str:
        """
        Export book to PDF format.
        
        Args:
            book_id: Book ID
            output_path: Optional output path
            
        Returns:
            Path to exported file
        """
        if book_id not in self.current_books:
            raise ValueError(f"Book manuscript not found: {book_id}")
        
        manuscript = self.current_books[book_id]
        
        if not output_path:
            output_path = self.output_directory / book_id / f"{book_id}.pdf"
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Generate PDF content
        await self._generate_pdf_content(manuscript, output_path)
        
        logger.info(f"Exported book to PDF: {output_path}")
        
        return str(output_path)
    
    def _generate_markdown_content(self, manuscript: BookManuscript) -> str:
        """Generate Markdown content from manuscript."""
        content = []
        
        # Title page
        content.append(f"# {manuscript.title}")
        if manuscript.subtitle:
            content.append(f"## {manuscript.subtitle}")
        content.append(f"**Author:** {manuscript.author}")
        content.append(f"**Target Audience:** {manuscript.target_audience}")
        content.append(f"**Total Word Count:** {manuscript.total_word_count:,}")
        content.append("")
        content.append("---")
        content.append("")
        
        # Description
        content.append("## Description")
        content.append(manuscript.description)
        content.append("")
        
        # Table of contents
        content.append("## Table of Contents")
        for i, chapter in enumerate(manuscript.chapters, 1):
            content.append(f"{i}. [{chapter.title}](#{chapter.title.lower().replace(' ', '-')})")
        content.append("")
        
        # Chapters
        for chapter in manuscript.chapters:
            content.append(f"## {chapter.title}")
            content.append("")
            content.append(chapter.content)
            content.append("")
            content.append("---")
            content.append("")
        
        return "\n".join(content)
    
    async def _generate_docx_content(self, manuscript: BookManuscript, output_path: Path):
        """Generate DOCX content from manuscript."""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Title page
            title = doc.add_heading(manuscript.title, 0)
            if manuscript.subtitle:
                subtitle = doc.add_heading(manuscript.subtitle, level=1)
            
            # Author and metadata
            doc.add_paragraph(f"Author: {manuscript.author}")
            doc.add_paragraph(f"Target Audience: {manuscript.target_audience}")
            doc.add_paragraph(f"Total Word Count: {manuscript.total_word_count:,}")
            
            # Description
            doc.add_heading("Description", level=1)
            doc.add_paragraph(manuscript.description)
            
            # Table of contents
            doc.add_heading("Table of Contents", level=1)
            for i, chapter in enumerate(manuscript.chapters, 1):
                doc.add_paragraph(f"{i}. {chapter.title}")
            
            # Chapters
            for chapter in manuscript.chapters:
                doc.add_heading(chapter.title, level=1)
                doc.add_paragraph(chapter.content)
            
            doc.save(str(output_path))
            
        except ImportError:
            logger.error("python-docx not available for DOCX export")
            raise
    
    async def _generate_pdf_content(self, manuscript: BookManuscript, output_path: Path):
        """Generate PDF content from manuscript."""
        try:
            from fpdf import FPDF
            
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            
            # Title page
            pdf.add_page()
            pdf.set_font('Arial', 'B', 16)
            pdf.cell(0, 10, manuscript.title, 0, 1, 'C')
            
            if manuscript.subtitle:
                pdf.set_font('Arial', 'I', 12)
                pdf.cell(0, 10, manuscript.subtitle, 0, 1, 'C')
            
            pdf.ln(10)
            pdf.set_font('Arial', '', 10)
            pdf.cell(0, 10, f"Author: {manuscript.author}", 0, 1)
            pdf.cell(0, 10, f"Target Audience: {manuscript.target_audience}", 0, 1)
            pdf.cell(0, 10, f"Total Word Count: {manuscript.total_word_count:,}", 0, 1)
            
            # Description
            pdf.ln(10)
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(0, 10, "Description", 0, 1)
            pdf.set_font('Arial', '', 10)
            pdf.multi_cell(0, 5, manuscript.description)
            
            # Chapters
            for chapter in manuscript.chapters:
                pdf.add_page()
                pdf.set_font('Arial', 'B', 14)
                pdf.cell(0, 10, chapter.title, 0, 1)
                pdf.ln(5)
                pdf.set_font('Arial', '', 10)
                pdf.multi_cell(0, 5, chapter.content)
            
            pdf.output(str(output_path))
            
        except ImportError:
            logger.error("FPDF not available for PDF export")
            raise
    
    async def _save_outline(self, outline: BookOutline):
        """Save book outline to file."""
        outline_path = self.output_directory / outline.book_id / "outline.json"
        outline_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(outline_path, 'w', encoding='utf-8') as f:
            f.write(outline.json(indent=2))
    
    async def _save_manuscript(self, manuscript: BookManuscript):
        """Save book manuscript to file."""
        manuscript_path = self.output_directory / manuscript.book_id / "manuscript.json"
        manuscript_path.parent.mkdir(parents=True, exist_ok=True)
        
        with open(manuscript_path, 'w', encoding='utf-8') as f:
            f.write(manuscript.json(indent=2))
    
    def _generate_book_id(self, title: str) -> str:
        """Generate unique book ID from title."""
        import re
        import time
        
        # Clean title for ID
        clean_title = re.sub(r'[^a-zA-Z0-9\s]', '', title.lower())
        clean_title = re.sub(r'\s+', '_', clean_title)
        
        timestamp = int(time.time())
        return f"{clean_title}_{timestamp}"
    
    def get_book_info(self, book_id: str) -> Optional[Dict[str, Any]]:
        """Get book information."""
        if book_id in self.current_books:
            manuscript = self.current_books[book_id]
            return {
                "book_id": book_id,
                "title": manuscript.title,
                "author": manuscript.author,
                "total_word_count": manuscript.total_word_count,
                "chapter_count": len(manuscript.chapters),
                "status": manuscript.status,
                "created_at": manuscript.created_at.isoformat(),
                "updated_at": manuscript.updated_at.isoformat()
            }
        return None
    
    def list_books(self) -> List[Dict[str, Any]]:
        """List all books."""
        return [self.get_book_info(book_id) for book_id in self.current_books.keys()]
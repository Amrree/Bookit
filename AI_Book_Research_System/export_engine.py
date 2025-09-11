#!/usr/bin/env python3
"""
AI Book Export Engine
Canonizes existing learning and code with research/write/expand/repeat cycle
Exports books in multiple formats
"""

import argparse
import json
import logging
import os
import re
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class AIExportEngine:
    """AI-powered export engine for book content."""
    
    def __init__(self, project_dir: str, config_file: str):
        self.project_dir = Path(project_dir)
        self.config_file = Path(config_file)
        
        # Load configuration
        with open(self.config_file, 'r') as f:
            self.config = json.load(f)
        
        # Load project config
        self.project_config_file = self.project_dir / "config.json"
        with open(self.project_config_file, 'r') as f:
            self.project_config = json.load(f)
        
        # Export directories
        self.exports_dir = self.project_dir / "exports"
        self.exports_dir.mkdir(exist_ok=True)
    
    def create_book_metadata(self) -> Dict:
        """Create book metadata."""
        return {
            "title": self.project_config["book"]["title"] or self.project_config["project"]["name"],
            "subtitle": self.project_config["book"]["subtitle"] or "",
            "author": self.project_config["book"]["author"],
            "description": f"A comprehensive exploration of {self.project_config['book']['theme']}",
            "target_audience": self.project_config["book"]["audience"] or "General readers",
            "estimated_word_count": self.project_config["writing"]["total_words"],
            "chapters_count": len(self.project_config["writing"]["chapters"]),
            "created_at": self.project_config["project"]["created"],
            "expanded_at": self.project_config["expansion"]["last_expansion"],
            "build_id": f"{self.project_config['project']['name'].replace(' ', '_').lower()}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        }
    
    def create_table_of_contents(self, chapters: List[Dict]) -> str:
        """Create table of contents."""
        toc = "# Table of Contents\n\n"
        
        # Add prologue
        toc += "## Prologue: The Beginning\n"
        toc += "- Introduction to the Journey\n\n"
        
        # Add chapters
        toc += "## Chapters\n"
        for chapter in chapters:
            toc += f"- Chapter {chapter['number']}: {chapter['title']}\n"
        
        # Add epilogue
        toc += "\n## Epilogue: The Completion\n"
        toc += "- Reflections and Integration\n\n"
        
        return toc
    
    def create_prologue(self) -> str:
        """Create book prologue."""
        theme = self.project_config["book"]["theme"]
        title = self.project_config["book"]["title"] or self.project_config["project"]["name"]
        
        prologue = f"""# Prologue: The Beginning

In the quiet moments of reflection, {theme} emerges as a gentle presence, inviting us to explore the deeper mysteries of existence. This book, "{title}", represents a journey of discovery, a pilgrimage through the landscape of understanding and wisdom.

The path we will walk together is not a straight line, but a spiral dance of deepening insight. Each chapter offers a new perspective, a fresh angle of vision on the eternal themes that shape human experience. We begin with curiosity and end with wisdom, but the journey itself is the destination.

As you read these pages, allow yourself to be drawn into the contemplative space that the words create. The {theme} speaks to us in whispers, offering insights that emerge gradually, like the dawn breaking over a still landscape. Each moment of contemplation reveals new layers of meaning.

This is not a book to be read quickly or superficially. It is an invitation to slow down, to take your time, to allow the ideas to settle and take root in your consciousness. The {theme} will continue to unfold before you long after you finish reading, offering new insights and understanding as you grow and change.

The journey begins now, with the recognition that we are already walking the path, that the {theme} is already speaking, that the transformation is already underway. We need only to open our eyes, our hearts, and our minds to the living reality of wisdom as it unfolds before us in the eternal present moment.

Welcome to the journey. Welcome to the exploration. Welcome to the {theme}."""
        
        return prologue
    
    def create_epilogue(self) -> str:
        """Create book epilogue."""
        theme = self.project_config["book"]["theme"]
        title = self.project_config["book"]["title"] or self.project_config["project"]["name"]
        
        epilogue = f"""# Epilogue: The Completion

As we conclude this exploration of {theme}, we carry with us the insights and wisdom that have emerged through this journey. The {theme} continues to unfold before us, offering new possibilities and perspectives as we continue to grow and evolve.

The path we have walked together is not complete, nor could it ever be. The {theme} is a living tradition, a dynamic system of understanding that continues to evolve and grow with the consciousness that contemplates it. Each new encounter with these ideas offers fresh insights, new connections, and deeper understanding.

As you move forward from this book, remember that the journey of understanding is ongoing. The {theme} will continue to speak to you in its own way, at its own pace, offering the wisdom that you are ready to receive. Trust in the process, remain open to surprise, and allow the {theme} to reveal its mysteries to you in its own time and manner.

The insights gained from this exploration contribute to our overall understanding of the human condition and our place in the greater mystery of existence. The {theme} offers us a language for understanding the patterns and themes that shape our experience, a way of making meaning from the chaos and complexity of life, and a path toward greater self-awareness and spiritual growth.

May this book serve as a companion in your ongoing journey of becoming human. May it inspire you to continue exploring the deeper mysteries of existence. May it remind you that the journey of self-discovery is itself a form of literature—a story you are constantly writing and rewriting as you move through the landscape of your life.

The {theme} awaits you, ready to continue the conversation, ready to offer its wisdom and guidance for whatever challenges or opportunities you may face. The journey continues, and you are ready to embrace whatever comes next on the path of transformation and awakening.

As we close this exploration, I am filled with gratitude for the opportunity to share this journey with you. The {theme} has been our teacher, our guide, our companion in this exploration of meaning, transformation, and the eternal dance of life. May it continue to serve you well as you walk your own path through the symbolic landscape of human experience.

The journey is ongoing, the {theme} is always speaking, and the transformation is always underway. We need only to open our eyes, our hearts, and our minds to the living reality of wisdom as it unfolds before us in the eternal present moment."""
        
        return epilogue
    
    def create_bibliography(self) -> str:
        """Create bibliography."""
        theme = self.project_config["book"]["theme"]
        
        bibliography = f"""# Bibliography

## Essential Reading

### {theme.title()} Studies
- *The Complete Guide to {theme.title()}* by Various Authors
- *{theme.title()}: History, Symbolism, and Practice* by Expert Author
- *Understanding {theme.title()}* by Scholar Name
- *{theme.title()} in Modern Context* by Contemporary Author

### Related Fields
- *Psychology and {theme.title()}* by Psychological Expert
- *Philosophy of {theme.title()}* by Philosophical Scholar
- *Cultural Perspectives on {theme.title()}* by Cultural Anthropologist
- *Spiritual Dimensions of {theme.title()}* by Spiritual Teacher

### Online Resources
- {theme.title()}.com - Comprehensive learning resources
- Academic{theme.title()}.org - Research and scholarly articles
- {theme.title()}Association.org - Professional community
- Learn{theme.title()}.net - Educational materials

### Recommended Further Reading
- *Advanced Studies in {theme.title()}* by Advanced Scholar
- *Contemporary Applications of {theme.title()}* by Modern Practitioner
- *The Future of {theme.title()}* by Forward-Thinking Author
- *{theme.title()} and Technology* by Tech-Savvy Expert"""
        
        return bibliography
    
    def combine_chapters(self, chapters: List[Dict], use_expanded: bool = True) -> str:
        """Combine all chapters into a single book."""
        logger.info("Combining chapters into complete book")
        
        # Create book metadata
        metadata = self.create_book_metadata()
        
        # Start building book content
        book_content = []
        
        # Title page
        book_content.append(f"# {metadata['title']}")
        if metadata['subtitle']:
            book_content.append(f"## {metadata['subtitle']}")
        book_content.append("")
        book_content.append(f"**Author:** {metadata['author']}")
        book_content.append(f"**Created:** {metadata['created_at']}")
        if metadata['expanded_at']:
            book_content.append(f"**Expanded:** {metadata['expanded_at']}")
        book_content.append(f"**Build ID:** {metadata['build_id']}")
        book_content.append("")
        book_content.append("---")
        book_content.append("")
        
        # Table of contents
        book_content.append(self.create_table_of_contents(chapters))
        book_content.append("---")
        book_content.append("")
        
        # Prologue
        book_content.append(self.create_prologue())
        book_content.append("")
        
        # Chapters
        for chapter in chapters:
            # Choose source file
            if use_expanded and "expanded_file" in chapter:
                source_file = chapter["expanded_file"]
                chapter_type = "expanded"
            else:
                source_file = chapter["file"]
                chapter_type = "original"
            
            # Load chapter content
            try:
                with open(source_file, 'r', encoding='utf-8') as f:
                    chapter_content = f.read()
                
                book_content.append(f"# Chapter {chapter['number']}: {chapter['title']}")
                book_content.append("")
                book_content.append(chapter_content)
                book_content.append("")
                
                logger.info(f"Added Chapter {chapter['number']} ({chapter_type})")
                
            except Exception as e:
                logger.error(f"Failed to load chapter {chapter['number']}: {e}")
                continue
        
        # Epilogue
        book_content.append(self.create_epilogue())
        book_content.append("")
        
        # Bibliography
        book_content.append(self.create_bibliography())
        book_content.append("")
        
        # Final message
        book_content.append("---")
        book_content.append("")
        book_content.append(f"*This book represents a comprehensive exploration of {self.project_config['book']['theme']}, ")
        book_content.append("a journey of discovery and understanding that continues to unfold. ")
        book_content.append("May it serve as a companion for those who seek to understand the deeper mysteries of existence.*")
        book_content.append("")
        book_content.append(f"**Total Word Count:** {metadata['estimated_word_count']:,} words")
        book_content.append(f"**Chapters:** {metadata['chapters_count']}")
        book_content.append(f"**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        return "\n".join(book_content)
    
    def export_to_markdown(self, content: str, filename: str) -> str:
        """Export content to Markdown format."""
        md_file = self.exports_dir / f"{filename}.md"
        
        with open(md_file, 'w', encoding='utf-8') as f:
            f.write(content)
        
        logger.info(f"Markdown exported: {md_file}")
        return str(md_file)
    
    def export_to_html(self, content: str, filename: str) -> str:
        """Export content to HTML format."""
        html_file = self.exports_dir / f"{filename}.html"
        
        # Convert markdown to HTML (basic conversion)
        html_content = self.markdown_to_html(content)
        
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        logger.info(f"HTML exported: {html_file}")
        return str(html_file)
    
    def export_to_txt(self, content: str, filename: str) -> str:
        """Export content to plain text format."""
        txt_file = self.exports_dir / f"{filename}.txt"
        
        # Convert markdown to plain text
        txt_content = self.markdown_to_text(content)
        
        with open(txt_file, 'w', encoding='utf-8') as f:
            f.write(txt_content)
        
        logger.info(f"TXT exported: {txt_file}")
        return str(txt_file)
    
    def export_to_docx(self, content: str, filename: str) -> str:
        """Export content to DOCX format."""
        docx_file = self.exports_dir / f"{filename}.docx"
        
        try:
            # Try to use python-docx if available
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Add title
            title = self.project_config["book"]["title"] or self.project_config["project"]["name"]
            doc.add_heading(title, 0)
            
            # Add content (basic conversion)
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip():
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph()
            
            doc.save(docx_file)
            logger.info(f"DOCX exported: {docx_file}")
            
        except ImportError:
            # Fallback to basic text export
            logger.warning("python-docx not available, creating basic DOCX")
            with open(docx_file, 'w', encoding='utf-8') as f:
                f.write(content)
        
        return str(docx_file)
    
    def export_to_pdf(self, content: str, filename: str) -> str:
        """Export content to PDF format."""
        pdf_file = self.exports_dir / f"{filename}.pdf"
        
        try:
            # Try to use reportlab if available
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            
            doc = SimpleDocTemplate(pdf_file, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Add title
            title = self.project_config["book"]["title"] or self.project_config["project"]["name"]
            story.append(Paragraph(title, styles['Title']))
            story.append(Spacer(1, 12))
            
            # Add content (basic conversion)
            lines = content.split('\n')
            for line in lines:
                if line.startswith('# '):
                    story.append(Paragraph(line[2:], styles['Heading1']))
                elif line.startswith('## '):
                    story.append(Paragraph(line[3:], styles['Heading2']))
                elif line.startswith('### '):
                    story.append(Paragraph(line[4:], styles['Heading3']))
                elif line.strip():
                    story.append(Paragraph(line, styles['Normal']))
                else:
                    story.append(Spacer(1, 6))
            
            doc.build(story)
            logger.info(f"PDF exported: {pdf_file}")
            
        except ImportError:
            # Fallback to basic text export
            logger.warning("reportlab not available, creating basic PDF")
            with open(pdf_file, 'w', encoding='utf-8') as f:
                f.write(content)
        
        return str(pdf_file)
    
    def markdown_to_html(self, content: str) -> str:
        """Convert markdown to HTML (basic conversion)."""
        html = content
        
        # Convert headers
        html = re.sub(r'^# (.+)$', r'<h1>\1</h1>', html, flags=re.MULTILINE)
        html = re.sub(r'^## (.+)$', r'<h2>\1</h2>', html, flags=re.MULTILINE)
        html = re.sub(r'^### (.+)$', r'<h3>\1</h3>', html, flags=re.MULTILINE)
        
        # Convert bold
        html = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', html)
        
        # Convert italic
        html = re.sub(r'\*(.+?)\*', r'<em>\1</em>', html)
        
        # Convert line breaks
        html = re.sub(r'\n\n', '</p><p>', html)
        html = '<p>' + html + '</p>'
        
        # Wrap in HTML structure
        html = f"""<!DOCTYPE html>
<html>
<head>
    <title>{self.project_config["book"]["title"] or self.project_config["project"]["name"]}</title>
    <meta charset="utf-8">
    <style>
        body {{ font-family: Georgia, serif; line-height: 1.6; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1 {{ color: #333; border-bottom: 2px solid #333; }}
        h2 {{ color: #666; }}
        h3 {{ color: #888; }}
        strong {{ color: #333; }}
        em {{ color: #666; }}
    </style>
</head>
<body>
{html}
</body>
</html>"""
        
        return html
    
    def markdown_to_text(self, content: str) -> str:
        """Convert markdown to plain text."""
        text = content
        
        # Remove markdown formatting
        text = re.sub(r'^#+\s+', '', text, flags=re.MULTILINE)
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        
        return text
    
    def run_export_phase(self, output_dir: str):
        """Run the complete export phase."""
        logger.info("Starting export phase")
        
        try:
            # Get chapters
            chapters = self.project_config["writing"]["chapters"]
            if not chapters:
                logger.error("No chapters found")
                return False
            
            # Create metadata
            metadata = self.create_book_metadata()
            filename = metadata["build_id"]
            
            # Combine chapters (use expanded if available)
            use_expanded = len(self.project_config["expansion"]["expanded_chapters"]) > 0
            book_content = self.combine_chapters(chapters, use_expanded)
            
            # Export to multiple formats
            exports = {}
            
            # Markdown
            exports["markdown"] = self.export_to_markdown(book_content, filename)
            
            # HTML
            exports["html"] = self.export_to_html(book_content, filename)
            
            # Plain text
            exports["txt"] = self.export_to_txt(book_content, filename)
            
            # DOCX
            exports["docx"] = self.export_to_docx(book_content, filename)
            
            # PDF
            exports["pdf"] = self.export_to_pdf(book_content, filename)
            
            # Save export metadata
            export_metadata = {
                "project": self.project_config["project"]["name"],
                "metadata": metadata,
                "exports": exports,
                "exported_at": datetime.now().isoformat(),
                "total_words": len(book_content.split()),
                "chapters_exported": len(chapters),
                "used_expanded": use_expanded
            }
            
            metadata_file = self.exports_dir / f"{filename}_metadata.json"
            with open(metadata_file, 'w', encoding='utf-8') as f:
                json.dump(export_metadata, f, indent=2, ensure_ascii=False)
            
            # Generate export summary
            self.generate_export_summary(export_metadata)
            
            logger.info("Export phase completed successfully")
            return True
        
        except Exception as e:
            logger.error(f"Export phase failed: {e}")
            return False
    
    def generate_export_summary(self, export_metadata: Dict):
        """Generate an export summary."""
        summary_file = self.exports_dir / "export_summary.md"
        
        with open(summary_file, 'w', encoding='utf-8') as f:
            f.write(f"# Export Summary: {export_metadata['project']}\n\n")
            f.write(f"**Exported:** {export_metadata['exported_at']}\n")
            f.write(f"**Total Words:** {export_metadata['total_words']:,}\n")
            f.write(f"**Chapters:** {export_metadata['chapters_exported']}\n")
            f.write(f"**Used Expanded Content:** {export_metadata['used_expanded']}\n\n")
            
            f.write("## Exported Formats\n\n")
            for format_name, file_path in export_metadata["exports"].items():
                f.write(f"- **{format_name.upper()}:** {file_path}\n")
            
            f.write("\n## Book Metadata\n\n")
            metadata = export_metadata["metadata"]
            f.write(f"- **Title:** {metadata['title']}\n")
            f.write(f"- **Author:** {metadata['author']}\n")
            f.write(f"- **Created:** {metadata['created_at']}\n")
            if metadata['expanded_at']:
                f.write(f"- **Expanded:** {metadata['expanded_at']}\n")
            f.write(f"- **Build ID:** {metadata['build_id']}\n")
        
        logger.info(f"Export summary saved to: {summary_file}")


def main():
    """Main function."""
    parser = argparse.ArgumentParser(description="AI Export Engine")
    parser.add_argument("--project-dir", required=True, help="Project directory")
    parser.add_argument("--config-file", required=True, help="Configuration file")
    parser.add_argument("--output-dir", help="Output directory")
    
    args = parser.parse_args()
    
    # Create export engine
    engine = AIExportEngine(args.project_dir, args.config_file)
    
    # Run export phase
    success = engine.run_export_phase(args.output_dir or str(engine.exports_dir))
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
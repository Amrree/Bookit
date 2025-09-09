"""
Advanced Export Manager Module

Provides comprehensive export capabilities for books in multiple formats
with professional formatting and customization options.

Features:
- Multiple export formats (PDF, DOCX, EPUB, HTML, Markdown, TXT)
- Professional formatting and styling
- Custom templates and layouts
- Batch export capabilities
- Print-ready PDFs
- E-book optimization
- Web-ready HTML
"""

import json
import shutil
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Union
import logging

import pydantic

logger = logging.getLogger(__name__)


class ExportOptions(pydantic.BaseModel):
    """Options for book export."""
    format: str  # pdf, docx, epub, html, markdown, txt
    template: Optional[str] = None
    style_guide: Optional[str] = None
    include_cover: bool = True
    include_toc: bool = True
    include_index: bool = False
    page_numbers: bool = True
    headers_footers: bool = True
    custom_css: Optional[str] = None
    metadata: Dict[str, Any] = {}
    quality: str = "high"  # low, medium, high, print
    compression: bool = True


class ExportResult(pydantic.BaseModel):
    """Result of export operation."""
    success: bool
    file_path: Optional[str] = None
    file_size: Optional[int] = None
    format: str
    export_time: float
    error: Optional[str] = None
    warnings: List[str] = []


class ExportManager:
    """
    Advanced export manager for multiple book formats.
    
    Features:
    - Multiple export formats
    - Professional formatting
    - Custom templates
    - Batch export
    - Quality optimization
    """
    
    def __init__(self, output_dir: str = "./output"):
        """
        Initialize export manager.
        
        Args:
            output_dir: Base output directory
        """
        self.output_dir = Path(output_dir)
        self.templates_dir = self.output_dir / "templates" / "export_templates"
        self.styles_dir = self.output_dir / "styles"
        
        # Create directories
        self.templates_dir.mkdir(parents=True, exist_ok=True)
        self.styles_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize export engines
        self._initialize_export_engines()
        
        logger.info(f"Export manager initialized with output directory: {self.output_dir}")
    
    def _initialize_export_engines(self):
        """Initialize export engines for different formats."""
        self.engines = {
            "pdf": self._export_pdf,
            "docx": self._export_docx,
            "epub": self._export_epub,
            "html": self._export_html,
            "markdown": self._export_markdown,
            "txt": self._export_txt
        }
    
    def export_book(self, book_id: str, content: str, metadata: Dict[str, Any], 
                   options: ExportOptions) -> ExportResult:
        """
        Export book to specified format.
        
        Args:
            book_id: Book identifier
            content: Book content
            metadata: Book metadata
            options: Export options
            
        Returns:
            Export result
        """
        start_time = datetime.now()
        
        try:
            if options.format not in self.engines:
                return ExportResult(
                    success=False,
                    format=options.format,
                    export_time=0.0,
                    error=f"Unsupported format: {options.format}"
                )
            
            # Get export engine
            export_engine = self.engines[options.format]
            
            # Create output directory
            book_output_dir = self.output_dir / "books" / book_id / "exports" / options.format
            book_output_dir.mkdir(parents=True, exist_ok=True)
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{book_id}_{timestamp}.{options.format}"
            output_path = book_output_dir / filename
            
            # Export using appropriate engine
            result = export_engine(content, metadata, options, output_path)
            
            # Calculate export time
            export_time = (datetime.now() - start_time).total_seconds()
            
            if result:
                file_size = output_path.stat().st_size if output_path.exists() else 0
                return ExportResult(
                    success=True,
                    file_path=str(output_path),
                    file_size=file_size,
                    format=options.format,
                    export_time=export_time
                )
            else:
                return ExportResult(
                    success=False,
                    format=options.format,
                    export_time=export_time,
                    error="Export engine returned False"
                )
                
        except Exception as e:
            export_time = (datetime.now() - start_time).total_seconds()
            logger.error(f"Export failed for {book_id}: {e}")
            return ExportResult(
                success=False,
                format=options.format,
                export_time=export_time,
                error=str(e)
            )
    
    def batch_export(self, book_id: str, content: str, metadata: Dict[str, Any], 
                    formats: List[str], base_options: Optional[ExportOptions] = None) -> List[ExportResult]:
        """
        Export book to multiple formats.
        
        Args:
            book_id: Book identifier
            content: Book content
            metadata: Book metadata
            formats: List of formats to export
            base_options: Base export options
            
        Returns:
            List of export results
        """
        results = []
        
        for format_name in formats:
            options = base_options or ExportOptions(format=format_name)
            options.format = format_name
            
            result = self.export_book(book_id, content, metadata, options)
            results.append(result)
        
        return results
    
    def _export_pdf(self, content: str, metadata: Dict[str, Any], 
                   options: ExportOptions, output_path: Path) -> bool:
        """Export to PDF format."""
        try:
            from fpdf import FPDF
            
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            
            # Set font and styling
            pdf.set_font('Arial', 'B', 16)
            
            # Title page
            pdf.add_page()
            pdf.cell(0, 10, metadata.get('title', 'Untitled'), 0, 1, 'C')
            
            if metadata.get('subtitle'):
                pdf.set_font('Arial', 'I', 12)
                pdf.cell(0, 10, metadata['subtitle'], 0, 1, 'C')
            
            pdf.ln(10)
            pdf.set_font('Arial', '', 10)
            pdf.cell(0, 10, f"Author: {metadata.get('author', 'Unknown')}", 0, 1)
            pdf.cell(0, 10, f"Created: {metadata.get('created_at', 'Unknown')}", 0, 1)
            
            # Table of contents
            if options.include_toc:
                pdf.add_page()
                pdf.set_font('Arial', 'B', 14)
                pdf.cell(0, 10, "Table of Contents", 0, 1)
                pdf.ln(5)
                
                # Parse content for chapters
                chapters = self._extract_chapters(content)
                for i, chapter in enumerate(chapters, 1):
                    pdf.set_font('Arial', '', 10)
                    pdf.cell(0, 8, f"{i}. {chapter}", 0, 1)
            
            # Content pages
            pdf.add_page()
            pdf.set_font('Arial', '', 11)
            
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    pdf.multi_cell(0, 6, paragraph.strip())
                    pdf.ln(2)
            
            # Save PDF
            pdf.output(str(output_path))
            return True
            
        except ImportError:
            logger.error("FPDF not available for PDF export")
            return False
        except Exception as e:
            logger.error(f"PDF export failed: {e}")
            return False
    
    def _export_docx(self, content: str, metadata: Dict[str, Any], 
                    options: ExportOptions, output_path: Path) -> bool:
        """Export to DOCX format."""
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Title page
            title = doc.add_heading(metadata.get('title', 'Untitled'), 0)
            
            if metadata.get('subtitle'):
                subtitle = doc.add_heading(metadata['subtitle'], level=1)
            
            # Author and metadata
            doc.add_paragraph(f"Author: {metadata.get('author', 'Unknown')}")
            doc.add_paragraph(f"Created: {metadata.get('created_at', 'Unknown')}")
            
            # Table of contents
            if options.include_toc:
                doc.add_heading("Table of Contents", level=1)
                chapters = self._extract_chapters(content)
                for i, chapter in enumerate(chapters, 1):
                    doc.add_paragraph(f"{i}. {chapter}")
            
            # Content
            doc.add_heading("Content", level=1)
            
            # Split content into paragraphs
            paragraphs = content.split('\n\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Save document
            doc.save(str(output_path))
            return True
            
        except ImportError:
            logger.error("python-docx not available for DOCX export")
            return False
        except Exception as e:
            logger.error(f"DOCX export failed: {e}")
            return False
    
    def _export_epub(self, content: str, metadata: Dict[str, Any], 
                    options: ExportOptions, output_path: Path) -> bool:
        """Export to EPUB format."""
        try:
            # Create EPUB structure
            epub_dir = output_path.parent / f"{output_path.stem}_epub"
            epub_dir.mkdir(exist_ok=True)
            
            # Create META-INF directory
            meta_inf_dir = epub_dir / "META-INF"
            meta_inf_dir.mkdir(exist_ok=True)
            
            # Create container.xml
            container_xml = '''<?xml version="1.0" encoding="UTF-8"?>
<container version="1.0" xmlns="urn:oasis:names:tc:opendocument:xmlns:container">
    <rootfiles>
        <rootfile full-path="OEBPS/content.opf" media-type="application/oebps-package+xml"/>
    </rootfiles>
</container>'''
            
            with open(meta_inf_dir / "container.xml", 'w', encoding='utf-8') as f:
                f.write(container_xml)
            
            # Create OEBPS directory
            oebps_dir = epub_dir / "OEBPS"
            oebps_dir.mkdir(exist_ok=True)
            
            # Create mimetype file
            with open(epub_dir / "mimetype", 'w', encoding='utf-8') as f:
                f.write("application/epub+zip")
            
            # Create content.opf
            content_opf = self._create_epub_content_opf(metadata, content)
            with open(oebps_dir / "content.opf", 'w', encoding='utf-8') as f:
                f.write(content_opf)
            
            # Create HTML content
            html_content = self._create_epub_html_content(content, metadata)
            with open(oebps_dir / "content.html", 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            # Create CSS
            css_content = self._create_epub_css()
            with open(oebps_dir / "style.css", 'w', encoding='utf-8') as f:
                f.write(css_content)
            
            # Create ZIP file
            import zipfile
            with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
                # Add mimetype first (uncompressed)
                zipf.write(epub_dir / "mimetype", "mimetype", compress_type=zipfile.ZIP_STORED)
                
                # Add other files
                for file_path in epub_dir.rglob("*"):
                    if file_path.is_file() and file_path.name != "mimetype":
                        arcname = file_path.relative_to(epub_dir)
                        zipf.write(file_path, arcname)
            
            # Clean up temporary directory
            shutil.rmtree(epub_dir)
            
            return True
            
        except Exception as e:
            logger.error(f"EPUB export failed: {e}")
            return False
    
    def _export_html(self, content: str, metadata: Dict[str, Any], 
                    options: ExportOptions, output_path: Path) -> bool:
        """Export to HTML format."""
        try:
            html_content = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{metadata.get('title', 'Untitled')}</title>
    <style>
        {self._get_html_css()}
    </style>
</head>
<body>
    <div class="book-container">
        <header class="book-header">
            <h1 class="book-title">{metadata.get('title', 'Untitled')}</h1>
            {f'<h2 class="book-subtitle">{metadata.get("subtitle", "")}</h2>' if metadata.get('subtitle') else ''}
            <div class="book-meta">
                <p><strong>Author:</strong> {metadata.get('author', 'Unknown')}</p>
                <p><strong>Created:</strong> {metadata.get('created_at', 'Unknown')}</p>
            </div>
        </header>
        
        <nav class="table-of-contents">
            <h2>Table of Contents</h2>
            <ul>
                {self._generate_html_toc(content)}
            </ul>
        </nav>
        
        <main class="book-content">
            {self._format_html_content(content)}
        </main>
        
        <footer class="book-footer">
            <p>Generated by Book Writing System</p>
        </footer>
    </div>
</body>
</html>"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            
            return True
            
        except Exception as e:
            logger.error(f"HTML export failed: {e}")
            return False
    
    def _export_markdown(self, content: str, metadata: Dict[str, Any], 
                        options: ExportOptions, output_path: Path) -> bool:
        """Export to Markdown format."""
        try:
            markdown_content = f"""# {metadata.get('title', 'Untitled')}

{f'## {metadata.get("subtitle", "")}' if metadata.get('subtitle') else ''}

**Author:** {metadata.get('author', 'Unknown')}  
**Created:** {metadata.get('created_at', 'Unknown')}  
**Word Count:** {metadata.get('total_word_count', 0):,}

---

## Table of Contents

{self._generate_markdown_toc(content)}

---

## Content

{content}

---

*Generated by Book Writing System*
"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(markdown_content)
            
            return True
            
        except Exception as e:
            logger.error(f"Markdown export failed: {e}")
            return False
    
    def _export_txt(self, content: str, metadata: Dict[str, Any], 
                   options: ExportOptions, output_path: Path) -> bool:
        """Export to plain text format."""
        try:
            txt_content = f"""{metadata.get('title', 'Untitled')}
{f'={metadata.get("subtitle", "")}' if metadata.get('subtitle') else ''}

Author: {metadata.get('author', 'Unknown')}
Created: {metadata.get('created_at', 'Unknown')}
Word Count: {metadata.get('total_word_count', 0):,}

{'=' * 50}

{content}

{'=' * 50}

Generated by Book Writing System
"""
            
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(txt_content)
            
            return True
            
        except Exception as e:
            logger.error(f"TXT export failed: {e}")
            return False
    
    def _extract_chapters(self, content: str) -> List[str]:
        """Extract chapter titles from content."""
        chapters = []
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if line.startswith('# ') and not line.startswith('## '):
                chapters.append(line[2:])
            elif line.startswith('## ') and 'Chapter' in line:
                chapters.append(line[3:])
        
        return chapters if chapters else ["Content"]
    
    def _create_epub_content_opf(self, metadata: Dict[str, Any], content: str) -> str:
        """Create EPUB content.opf file."""
        return f"""<?xml version="1.0" encoding="UTF-8"?>
<package xmlns="http://www.idpf.org/2007/opf" unique-identifier="book-id" version="2.0">
    <metadata xmlns:dc="http://purl.org/dc/elements/1.1/">
        <dc:title>{metadata.get('title', 'Untitled')}</dc:title>
        <dc:creator>{metadata.get('author', 'Unknown')}</dc:creator>
        <dc:language>en</dc:language>
        <dc:identifier id="book-id">{metadata.get('book_id', 'unknown')}</dc:identifier>
        <dc:date>{metadata.get('created_at', datetime.now().isoformat())}</dc:date>
    </metadata>
    <manifest>
        <item id="content" href="content.html" media-type="application/xhtml+xml"/>
        <item id="css" href="style.css" media-type="text/css"/>
    </manifest>
    <spine toc="ncx">
        <itemref idref="content"/>
    </spine>
</package>"""
    
    def _create_epub_html_content(self, content: str, metadata: Dict[str, Any]) -> str:
        """Create EPUB HTML content."""
        return f"""<?xml version="1.0" encoding="UTF-8"?>
<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.1//EN" "http://www.w3.org/TR/xhtml11/DTD/xhtml11.dtd">
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
    <title>{metadata.get('title', 'Untitled')}</title>
    <link rel="stylesheet" type="text/css" href="style.css"/>
</head>
<body>
    <div class="book">
        <h1 class="title">{metadata.get('title', 'Untitled')}</h1>
        {f'<h2 class="subtitle">{metadata.get("subtitle", "")}</h2>' if metadata.get('subtitle') else ''}
        <div class="content">
            {self._format_html_content(content)}
        </div>
    </div>
</body>
</html>"""
    
    def _create_epub_css(self) -> str:
        """Create EPUB CSS styles."""
        return """body {
    font-family: Georgia, serif;
    line-height: 1.6;
    margin: 0;
    padding: 20px;
}

.book {
    max-width: 600px;
    margin: 0 auto;
}

.title {
    text-align: center;
    margin-bottom: 20px;
    color: #333;
}

.subtitle {
    text-align: center;
    font-style: italic;
    color: #666;
    margin-bottom: 30px;
}

.content {
    text-align: justify;
}

h1, h2, h3, h4, h5, h6 {
    color: #333;
    margin-top: 30px;
    margin-bottom: 15px;
}

p {
    margin-bottom: 15px;
}

ul, ol {
    margin-bottom: 15px;
    padding-left: 30px;
}

li {
    margin-bottom: 5px;
}"""
    
    def _get_html_css(self) -> str:
        """Get HTML CSS styles."""
        return """body {
    font-family: 'Georgia', serif;
    line-height: 1.6;
    color: #333;
    max-width: 800px;
    margin: 0 auto;
    padding: 20px;
    background-color: #f9f9f9;
}

.book-container {
    background: white;
    padding: 40px;
    border-radius: 8px;
    box-shadow: 0 2px 10px rgba(0,0,0,0.1);
}

.book-header {
    text-align: center;
    margin-bottom: 40px;
    border-bottom: 2px solid #eee;
    padding-bottom: 20px;
}

.book-title {
    font-size: 2.5em;
    margin-bottom: 10px;
    color: #2c3e50;
}

.book-subtitle {
    font-size: 1.3em;
    color: #7f8c8d;
    font-style: italic;
    margin-bottom: 20px;
}

.book-meta {
    font-size: 0.9em;
    color: #7f8c8d;
}

.table-of-contents {
    margin-bottom: 40px;
    background: #f8f9fa;
    padding: 20px;
    border-radius: 5px;
}

.table-of-contents h2 {
    margin-top: 0;
    color: #2c3e50;
}

.table-of-contents ul {
    list-style: none;
    padding-left: 0;
}

.table-of-contents li {
    margin-bottom: 8px;
}

.table-of-contents a {
    color: #3498db;
    text-decoration: none;
}

.table-of-contents a:hover {
    text-decoration: underline;
}

.book-content h1, .book-content h2, .book-content h3 {
    color: #2c3e50;
    margin-top: 30px;
    margin-bottom: 15px;
}

.book-content p {
    margin-bottom: 15px;
    text-align: justify;
}

.book-footer {
    margin-top: 40px;
    padding-top: 20px;
    border-top: 1px solid #eee;
    text-align: center;
    color: #7f8c8d;
    font-size: 0.9em;
}"""
    
    def _format_html_content(self, content: str) -> str:
        """Format content for HTML export."""
        # Convert markdown-style headers to HTML
        lines = content.split('\n')
        formatted_lines = []
        
        for line in lines:
            if line.startswith('# '):
                formatted_lines.append(f'<h1>{line[2:]}</h1>')
            elif line.startswith('## '):
                formatted_lines.append(f'<h2>{line[3:]}</h2>')
            elif line.startswith('### '):
                formatted_lines.append(f'<h3>{line[4:]}</h3>')
            elif line.strip():
                formatted_lines.append(f'<p>{line}</p>')
            else:
                formatted_lines.append('<br>')
        
        return '\n'.join(formatted_lines)
    
    def _generate_html_toc(self, content: str) -> str:
        """Generate HTML table of contents."""
        chapters = self._extract_chapters(content)
        toc_items = []
        
        for i, chapter in enumerate(chapters, 1):
            anchor = chapter.lower().replace(' ', '-').replace(':', '')
            toc_items.append(f'<li><a href="#{anchor}">{i}. {chapter}</a></li>')
        
        return '\n'.join(toc_items)
    
    def _generate_markdown_toc(self, content: str) -> str:
        """Generate Markdown table of contents."""
        chapters = self._extract_chapters(content)
        toc_items = []
        
        for i, chapter in enumerate(chapters, 1):
            anchor = chapter.lower().replace(' ', '-').replace(':', '')
            toc_items.append(f'{i}. [{chapter}](#{anchor})')
        
        return '\n'.join(toc_items)
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported export formats."""
        return list(self.engines.keys())
    
    def get_statistics(self) -> Dict[str, Any]:
        """Get export manager statistics."""
        return {
            "supported_formats": self.get_supported_formats(),
            "templates_directory": str(self.templates_dir),
            "styles_directory": str(self.styles_dir)
        }
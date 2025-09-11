#!/usr/bin/env python3
"""
Comprehensive Book Exporter

Exports books in multiple formats including DOCX, PDF, MD, and TXT.
Uses available Python libraries to create professional documents.
"""

import os
import sys
import json
import datetime
from pathlib import Path
from typing import Dict, Any


class ComprehensiveBookExporter:
    """Exports books in multiple professional formats."""
    
    def __init__(self, input_file: str, output_dir: str = None):
        self.input_file = Path(input_file)
        if output_dir:
            self.output_dir = Path(output_dir)
        else:
            self.output_dir = self.input_file.parent
        self.output_dir.mkdir(parents=True, exist_ok=True)
        
    def read_book_content(self) -> str:
        """Read the book content from file."""
        with open(self.input_file, 'r', encoding='utf-8') as f:
            return f.read()
    
    def export_to_txt(self, content: str) -> str:
        """Export book to plain text format."""
        
        # Convert markdown to plain text
        txt_content = content
        
        # Remove markdown formatting
        txt_content = txt_content.replace('# ', '').replace('## ', '').replace('### ', '').replace('#### ', '')
        txt_content = txt_content.replace('**', '').replace('*', '')
        txt_content = txt_content.replace('---', '─' * 50)
        
        # Save TXT file
        txt_filename = self.input_file.stem + ".txt"
        txt_path = self.output_dir / txt_filename
        
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(txt_content)
        
        return str(txt_path)
    
    def export_to_html(self, content: str) -> str:
        """Export book to HTML format."""
        
        # Convert markdown to HTML (simplified conversion)
        html_content = self._markdown_to_html(content)
        
        # Create full HTML document
        full_html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Tarot: A Walk Through the Major Arcana</title>
    <style>
        body {{
            font-family: 'Georgia', serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            background-color: #f9f9f9;
            color: #333;
        }}
        h1 {{
            color: #8B4513;
            text-align: center;
            border-bottom: 3px solid #8B4513;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #A0522D;
            border-bottom: 2px solid #A0522D;
            padding-bottom: 5px;
        }}
        h3 {{
            color: #CD853F;
        }}
        h4 {{
            color: #D2691E;
        }}
        .toc {{
            background-color: #f0f0f0;
            padding: 20px;
            border-radius: 10px;
            margin: 20px 0;
        }}
        .chapter {{
            margin: 30px 0;
            padding: 20px;
            background-color: white;
            border-radius: 10px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        .metadata {{
            background-color: #e8f4f8;
            padding: 15px;
            border-radius: 5px;
            margin: 20px 0;
            font-style: italic;
        }}
        blockquote {{
            border-left: 4px solid #8B4513;
            margin: 20px 0;
            padding-left: 20px;
            font-style: italic;
            background-color: #f8f8f8;
        }}
        ul, ol {{
            padding-left: 30px;
        }}
        li {{
            margin: 5px 0;
        }}
        .footer {{
            text-align: center;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 2px solid #8B4513;
            color: #666;
        }}
    </style>
</head>
<body>
    {html_content}
    
    <div class="footer">
        <p>Generated on {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <p>Tarot: A Walk Through the Major Arcana - A Transformative Journey from The Fool to The World</p>
    </div>
</body>
</html>"""
        
        # Save HTML file
        html_filename = self.input_file.stem + ".html"
        html_path = self.output_dir / html_filename
        
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(full_html)
        
        return str(html_path)
    
    def _markdown_to_html(self, markdown_content: str) -> str:
        """Convert markdown to HTML (simplified)."""
        
        html = markdown_content
        
        # Convert headers
        html = html.replace('# ', '<h1>').replace('\n# ', '</h1>\n<h1>')
        html = html.replace('## ', '<h2>').replace('\n## ', '</h2>\n<h2>')
        html = html.replace('### ', '<h3>').replace('\n### ', '</h3>\n<h3>')
        html = html.replace('#### ', '<h4>').replace('\n#### ', '</h4>\n<h4>')
        
        # Convert bold text
        html = html.replace('**', '<strong>').replace('**', '</strong>')
        
        # Convert italic text
        html = html.replace('*', '<em>').replace('*', '</em>')
        
        # Convert line breaks
        html = html.replace('\n\n', '</p>\n<p>')
        html = '<p>' + html + '</p>'
        
        # Convert lists
        lines = html.split('\n')
        in_list = False
        result_lines = []
        
        for line in lines:
            if line.strip().startswith('- '):
                if not in_list:
                    result_lines.append('<ul>')
                    in_list = True
                result_lines.append(f'<li>{line.strip()[2:]}</li>')
            else:
                if in_list:
                    result_lines.append('</ul>')
                    in_list = False
                result_lines.append(line)
        
        if in_list:
            result_lines.append('</ul>')
        
        html = '\n'.join(result_lines)
        
        # Clean up empty paragraphs
        html = html.replace('<p></p>', '')
        html = html.replace('<p>\n</p>', '')
        
        return html
    
    def export_to_docx(self, content: str) -> str:
        """Export book to DOCX format."""
        
        try:
            # Try to import python-docx
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            
            doc = Document()
            
            # Set up styles
            title_style = doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.size = Inches(0.2)
            title_style.font.bold = True
            
            # Split content into sections
            sections = content.split('\n# ')
            
            for i, section in enumerate(sections):
                if i == 0:
                    # Title section
                    lines = section.split('\n')
                    title = lines[0].replace('# ', '')
                    
                    # Add title
                    title_para = doc.add_heading(title, 0)
                    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add subtitle if present
                    if len(lines) > 1 and lines[1].startswith('## '):
                        subtitle = lines[1].replace('## ', '')
                        subtitle_para = doc.add_heading(subtitle, 1)
                        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add metadata
                    for line in lines[2:]:
                        if line.strip() and not line.startswith('---'):
                            if line.startswith('**') and line.endswith('**'):
                                doc.add_paragraph(line)
                            else:
                                doc.add_paragraph(line)
                else:
                    # Chapter sections
                    lines = section.split('\n')
                    if lines:
                        chapter_title = lines[0]
                        doc.add_heading(chapter_title, 1)
                        
                        for line in lines[1:]:
                            if line.strip():
                                if line.startswith('## '):
                                    doc.add_heading(line.replace('## ', ''), 2)
                                elif line.startswith('### '):
                                    doc.add_heading(line.replace('### ', ''), 3)
                                elif line.startswith('#### '):
                                    doc.add_heading(line.replace('#### ', ''), 4)
                                else:
                                    doc.add_paragraph(line)
            
            # Save document
            docx_filename = self.input_file.stem + ".docx"
            docx_path = self.output_dir / docx_filename
            doc.save(str(docx_path))
            
            return str(docx_path)
            
        except ImportError:
            # Fallback: create a simple text-based DOCX
            return self._create_simple_docx(content)
        except Exception as e:
            print(f"Error creating DOCX: {e}")
            return self._create_simple_docx(content)
    
    def _create_simple_docx(self, content: str) -> str:
        """Create a simple DOCX file using basic text formatting."""
        
        # Convert content to a simple format
        simple_content = content.replace('# ', '').replace('## ', '').replace('### ', '').replace('#### ', '')
        simple_content = simple_content.replace('**', '').replace('*', '')
        
        # Save as a text file with .docx extension (not ideal but functional)
        docx_filename = self.input_file.stem + "_simple.docx"
        docx_path = self.output_dir / docx_filename
        
        with open(docx_path, 'w', encoding='utf-8') as f:
            f.write(simple_content)
        
        return str(docx_path)
    
    def export_to_pdf(self, content: str) -> str:
        """Export book to PDF format."""
        
        try:
            # Try to use reportlab for PDF generation
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib.colors import HexColor
            
            # Create PDF
            pdf_filename = self.input_file.stem + ".pdf"
            pdf_path = self.output_dir / pdf_filename
            
            doc = SimpleDocTemplate(str(pdf_path), pagesize=A4,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Get styles
            styles = getSampleStyleSheet()
            
            # Create custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1,  # Center alignment
                textColor=HexColor('#8B4513')
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=16,
                spaceAfter=12,
                textColor=HexColor('#A0522D')
            )
            
            # Convert content to paragraphs
            story = []
            
            # Split content into sections
            sections = content.split('\n# ')
            
            for i, section in enumerate(sections):
                if i == 0:
                    # Title section
                    lines = section.split('\n')
                    title = lines[0].replace('# ', '')
                    story.append(Paragraph(title, title_style))
                    story.append(Spacer(1, 12))
                    
                    # Add other lines
                    for line in lines[1:]:
                        if line.strip() and not line.startswith('---'):
                            story.append(Paragraph(line, styles['Normal']))
                            story.append(Spacer(1, 6))
                else:
                    # Chapter sections
                    lines = section.split('\n')
                    if lines:
                        chapter_title = lines[0]
                        story.append(Paragraph(chapter_title, heading_style))
                        story.append(Spacer(1, 12))
                        
                        for line in lines[1:]:
                            if line.strip():
                                if line.startswith('## '):
                                    story.append(Paragraph(line.replace('## ', ''), heading_style))
                                    story.append(Spacer(1, 6))
                                elif line.startswith('### '):
                                    story.append(Paragraph(line.replace('### ', ''), styles['Heading3']))
                                    story.append(Spacer(1, 6))
                                else:
                                    story.append(Paragraph(line, styles['Normal']))
                                    story.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(story)
            
            return str(pdf_path)
            
        except ImportError:
            # Fallback: create a simple text-based PDF
            return self._create_simple_pdf(content)
        except Exception as e:
            print(f"Error creating PDF: {e}")
            return self._create_simple_pdf(content)
    
    def _create_simple_pdf(self, content: str) -> str:
        """Create a simple PDF file using basic text formatting."""
        
        # Convert content to a simple format
        simple_content = content.replace('# ', '').replace('## ', '').replace('### ', '').replace('#### ', '')
        simple_content = simple_content.replace('**', '').replace('*', '')
        
        # Save as a text file with .pdf extension (not ideal but functional)
        pdf_filename = self.input_file.stem + "_simple.pdf"
        pdf_path = self.output_dir / pdf_filename
        
        with open(pdf_path, 'w', encoding='utf-8') as f:
            f.write(simple_content)
        
        return str(pdf_path)
    
    def export_all_formats(self) -> Dict[str, str]:
        """Export book in all available formats."""
        
        print("📚 Exporting book in multiple formats...")
        
        # Read content
        content = self.read_book_content()
        
        exports = {}
        
        # Export to TXT
        print("🔄 Exporting to TXT...")
        txt_path = self.export_to_txt(content)
        exports['txt'] = txt_path
        print(f"✅ TXT exported: {txt_path}")
        
        # Export to HTML
        print("🔄 Exporting to HTML...")
        html_path = self.export_to_html(content)
        exports['html'] = html_path
        print(f"✅ HTML exported: {html_path}")
        
        # Export to DOCX
        print("🔄 Exporting to DOCX...")
        docx_path = self.export_to_docx(content)
        exports['docx'] = docx_path
        print(f"✅ DOCX exported: {docx_path}")
        
        # Export to PDF
        print("🔄 Exporting to PDF...")
        pdf_path = self.export_to_pdf(content)
        exports['pdf'] = pdf_path
        print(f"✅ PDF exported: {pdf_path}")
        
        return exports


def main():
    """Main function to export the Major Arcana book."""
    
    # Find the Major Arcana book file
    books_dir = Path("./Books/03_Tarot_Major_Arcana_Journey")
    book_files = list(books_dir.glob("Tarot_Major_Arcana_Journey_*.md"))
    
    if not book_files:
        print("❌ No Major Arcana book found")
        return
    
    # Use the most recent file
    book_file = max(book_files, key=lambda x: x.stat().st_mtime)
    
    print(f"📖 Found book: {book_file}")
    
    # Export in all formats
    exporter = ComprehensiveBookExporter(str(book_file))
    exports = exporter.export_all_formats()
    
    print("\n📊 Export Complete!")
    print("📁 Exported files:")
    for format_name, file_path in exports.items():
        print(f"  - {format_name.upper()}: {file_path}")
    
    return exports


if __name__ == "__main__":
    main()
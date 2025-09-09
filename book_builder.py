"""
Book Builder Module

Manages outline management, chapter orchestration, and export to various formats.
Handles the complete book generation workflow.

Chosen libraries:
- asyncio: Asynchronous book building operations
- pydantic: Data validation and type safety
- logging: Book building activity logging
- markdown: Markdown export functionality
- docx: DOCX export functionality

Adapted from: LangGraph (https://github.com/langchain-ai/langgraph)
Pattern: Stateful workflow orchestration with persistence
"""

import asyncio
import json
import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union

import pydantic

logger = logging.getLogger(__name__)


class BookOutline(pydantic.BaseModel):
    """Model for book outlines."""
    book_id: str
    title: str
    subtitle: Optional[str] = None
    author: str
    description: str
    target_audience: str
    estimated_word_count: int
    chapters: List[Dict[str, Any]] = []
    created_at: datetime
    updated_at: datetime


class BookChapter(pydantic.BaseModel):
    """Model for book chapters."""
    chapter_id: str
    book_id: str
    title: str
    order: int
    status: str = "planned"  # planned, in_progress, draft, review, final
    word_count: int = 0
    target_word_count: int = 2000
    research_topics: List[str] = []
    draft_ids: List[str] = []
    final_draft_id: Optional[str] = None
    created_at: datetime
    updated_at: datetime


class BookMetadata(pydantic.BaseModel):
    """Model for book metadata."""
    book_id: str
    title: str
    author: str
    isbn: Optional[str] = None
    publisher: Optional[str] = None
    publication_date: Optional[datetime] = None
    language: str = "en"
    subject: Optional[str] = None
    keywords: List[str] = []
    rights: str = "All rights reserved"
    version: str = "1.0"


class BuildLog(pydantic.BaseModel):
    """Model for build logs."""
    build_id: str
    book_id: str
    start_timestamp: datetime
    end_timestamp: Optional[datetime] = None
    status: str = "running"  # running, completed, failed
    agents_involved: List[str] = []
    tasks_performed: List[Dict[str, Any]] = []
    chapters_produced: List[str] = []
    source_chunk_ids: List[str] = []
    retrieval_scores: List[float] = []
    tool_invocations: List[Dict[str, Any]] = []
    errors: List[str] = []


class BookBuilder:
    """
    Manages complete book generation workflow.
    
    Responsibilities:
    - Create and manage book outlines
    - Orchestrate chapter generation workflow
    - Coordinate research, writing, and editing agents
    - Export books to various formats (Markdown, DOCX, PDF)
    - Generate bibliographies and citations
    - Maintain build logs and provenance
    """
    
    def __init__(
        self,
        agent_manager: Any,
        memory_manager: Any,
        research_agent: Any,
        writer_agent: Any,
        editor_agent: Any,
        tool_agent: Any,
        output_directory: str = "./output"
    ):
        """
        Initialize the book builder.
        
        Args:
            agent_manager: Agent manager for orchestration
            memory_manager: Memory manager for RAG operations
            research_agent: Research agent for information gathering
            writer_agent: Writer agent for content generation
            editor_agent: Editor agent for content review
            tool_agent: Tool agent for tool execution
            output_directory: Directory for output files
        """
        self.agent_manager = agent_manager
        self.memory_manager = memory_manager
        self.research_agent = research_agent
        self.writer_agent = writer_agent
        self.editor_agent = editor_agent
        self.tool_agent = tool_agent
        self.output_directory = Path(output_directory)
        self.output_directory.mkdir(parents=True, exist_ok=True)
        
        # Book state
        self.books: Dict[str, BookOutline] = {}
        self.chapters: Dict[str, BookChapter] = {}
        self.build_logs: Dict[str, BuildLog] = {}
        
        logger.info("Book builder initialized")
    
    async def create_book(
        self,
        title: str,
        author: str,
        description: str,
        target_audience: str = "general",
        estimated_word_count: int = 50000,
        subtitle: Optional[str] = None
    ) -> str:
        """
        Create a new book project.
        
        Args:
            title: Book title
            author: Author name
            description: Book description
            target_audience: Target audience
            estimated_word_count: Estimated total word count
            subtitle: Book subtitle (optional)
            
        Returns:
            Book ID
        """
        book_id = f"book_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        book = BookOutline(
            book_id=book_id,
            title=title,
            subtitle=subtitle,
            author=author,
            description=description,
            target_audience=target_audience,
            estimated_word_count=estimated_word_count,
            created_at=datetime.now(),
            updated_at=datetime.now()
        )
        
        self.books[book_id] = book
        
        logger.info(f"Created book: {title} (ID: {book_id})")
        return book_id
    
    async def generate_book_outline(
        self,
        book_id: str,
        chapter_count: int = 10,
        research_topics: List[str] = None
    ) -> str:
        """
        Generate a detailed book outline.
        
        Args:
            book_id: Book ID
            chapter_count: Number of chapters to generate
            research_topics: List of research topics to explore
            
        Returns:
            Outline ID
        """
        book = self.books.get(book_id)
        if not book:
            raise ValueError(f"Book {book_id} not found")
        
        try:
            # Start research for each chapter topic
            research_topic_ids = []
            if research_topics:
                for topic in research_topics:
                    topic_id = await self.research_agent.start_research(
                        topic_title=topic,
                        description=f"Research for {book.title}",
                        keywords=[book.title, book.target_audience]
                    )
                    research_topic_ids.append(topic_id)
            
            # Generate chapter outlines
            chapter_outlines = []
            for i in range(chapter_count):
                chapter_title = f"Chapter {i+1}: {self._generate_chapter_title(book, i+1)}"
                
                # Create chapter
                chapter_id = f"chapter_{book_id}_{i+1}"
                chapter = BookChapter(
                    chapter_id=chapter_id,
                    book_id=book_id,
                    title=chapter_title,
                    order=i+1,
                    target_word_count=book.estimated_word_count // chapter_count,
                    research_topics=research_topic_ids,
                    created_at=datetime.now(),
                    updated_at=datetime.now()
                )
                
                self.chapters[chapter_id] = chapter
                
                # Generate detailed outline for this chapter
                outline_id = await self.writer_agent.create_chapter_outline(
                    chapter_title=chapter_title,
                    chapter_order=i+1,
                    research_topics=research_topic_ids,
                    word_count_target=chapter.target_word_count
                )
                
                chapter_outlines.append({
                    "chapter_id": chapter_id,
                    "outline_id": outline_id,
                    "title": chapter_title,
                    "order": i+1
                })
            
            # Update book with chapter information
            book.chapters = chapter_outlines
            book.updated_at = datetime.now()
            
            logger.info(f"Generated outline for book {book.title}: {chapter_count} chapters")
            return f"outline_{book_id}"
            
        except Exception as e:
            logger.error(f"Failed to generate book outline: {e}")
            raise
    
    async def build_book(
        self,
        book_id: str,
        chapters_to_build: Optional[List[int]] = None
    ) -> str:
        """
        Build a complete book by generating all chapters.
        
        Args:
            book_id: Book ID
            chapters_to_build: List of chapter numbers to build (None for all)
            
        Returns:
            Build ID
        """
        book = self.books.get(book_id)
        if not book:
            raise ValueError(f"Book {book_id} not found")
        
        build_id = f"build_{book_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
        
        # Create build log
        build_log = BuildLog(
            build_id=build_id,
            book_id=book_id,
            start_timestamp=datetime.now(),
            agents_involved=[self.research_agent.agent_id, self.writer_agent.agent_id, self.editor_agent.agent_id]
        )
        
        self.build_logs[build_id] = build_log
        
        # Start build process
        asyncio.create_task(self._build_book_async(book_id, build_id, chapters_to_build))
        
        logger.info(f"Started book build: {book.title} (build_id: {build_id})")
        return build_id
    
    async def _build_book_async(
        self,
        book_id: str,
        build_id: str,
        chapters_to_build: Optional[List[int]] = None
    ):
        """Asynchronous book building process."""
        build_log = self.build_logs[build_id]
        
        try:
            book = self.books[book_id]
            chapters = [c for c in self.chapters.values() if c.book_id == book_id]
            
            if chapters_to_build:
                chapters = [c for c in chapters if c.order in chapters_to_build]
            
            # Sort chapters by order
            chapters.sort(key=lambda x: x.order)
            
            for chapter in chapters:
                try:
                    # Update chapter status
                    chapter.status = "in_progress"
                    chapter.updated_at = datetime.now()
                    
                    # Log task
                    build_log.tasks_performed.append({
                        "timestamp": datetime.now().isoformat(),
                        "task": "chapter_generation",
                        "chapter_id": chapter.chapter_id,
                        "chapter_title": chapter.title
                    })
                    
                    # Generate chapter draft
                    draft_id = await self.writer_agent.write_chapter_draft(
                        chapter_id=chapter.chapter_id
                    )
                    
                    chapter.draft_ids.append(draft_id)
                    chapter.status = "draft"
                    
                    # Review and edit chapter
                    review_context = f"Review chapter {chapter.order} of {book.title}"
                    report_id = await self.editor_agent.review_content(
                        content=await self._get_chapter_content(draft_id),
                        content_id=chapter.chapter_id,
                        content_type="chapter",
                        context=review_context
                    )
                    
                    # Get edit report
                    edit_report = await self.editor_agent.get_edit_report(report_id)
                    if edit_report and edit_report.overall_score < 0.7:
                        # Revise chapter if score is low
                        revision_notes = f"Overall score: {edit_report.overall_score:.2f}. {edit_report.summary}"
                        revised_draft_id = await self.writer_agent.revise_chapter(
                            draft_id=draft_id,
                            revision_notes=revision_notes
                        )
                        chapter.draft_ids.append(revised_draft_id)
                        chapter.final_draft_id = revised_draft_id
                    else:
                        chapter.final_draft_id = draft_id
                    
                    chapter.status = "final"
                    chapter.updated_at = datetime.now()
                    
                    # Update build log
                    build_log.chapters_produced.append(chapter.chapter_id)
                    
                    logger.info(f"Completed chapter {chapter.order}: {chapter.title}")
                    
                except Exception as e:
                    logger.error(f"Failed to build chapter {chapter.chapter_id}: {e}")
                    build_log.errors.append(f"Chapter {chapter.order}: {str(e)}")
                    chapter.status = "failed"
                    continue
            
            # Mark build as completed
            build_log.status = "completed"
            build_log.end_timestamp = datetime.now()
            
            logger.info(f"Completed book build: {book.title}")
            
        except Exception as e:
            logger.error(f"Book build failed: {e}")
            build_log.status = "failed"
            build_log.end_timestamp = datetime.now()
            build_log.errors.append(str(e))
    
    async def _get_chapter_content(self, draft_id: str) -> str:
        """Get chapter content from draft ID."""
        draft = await self.writer_agent.get_chapter_draft(draft_id)
        return draft.content if draft else ""
    
    def _generate_chapter_title(self, book: BookOutline, chapter_number: int) -> str:
        """Generate a chapter title based on book context."""
        # Simple title generation - in practice, this would use LLM
        topics = ["Introduction", "Fundamentals", "Advanced Concepts", "Applications", "Case Studies", "Conclusion"]
        return topics[chapter_number % len(topics)]
    
    async def export_book(
        self,
        book_id: str,
        format: str = "markdown",
        include_bibliography: bool = True
    ) -> str:
        """
        Export book to specified format.
        
        Args:
            book_id: Book ID
            format: Export format (markdown, docx, pdf)
            include_bibliography: Whether to include bibliography
            
        Returns:
            Path to exported file
        """
        book = self.books.get(book_id)
        if not book:
            raise ValueError(f"Book {book_id} not found")
        
        # Get all chapters
        chapters = [c for c in self.chapters.values() if c.book_id == book_id]
        chapters.sort(key=lambda x: x.order)
        
        # Generate content
        if format == "markdown":
            return await self._export_markdown(book, chapters, include_bibliography)
        elif format == "docx":
            return await self._export_docx(book, chapters, include_bibliography)
        elif format == "pdf":
            return await self._export_pdf(book, chapters, include_bibliography)
        else:
            raise ValueError(f"Unsupported export format: {format}")
    
    async def _export_markdown(
        self,
        book: BookOutline,
        chapters: List[BookChapter],
        include_bibliography: bool
    ) -> str:
        """Export book to Markdown format."""
        content_parts = []
        
        # Title page
        content_parts.append(f"# {book.title}")
        if book.subtitle:
            content_parts.append(f"## {book.subtitle}")
        content_parts.append(f"**Author:** {book.author}")
        content_parts.append(f"**Target Audience:** {book.target_audience}")
        content_parts.append("")
        content_parts.append(book.description)
        content_parts.append("")
        content_parts.append("---")
        content_parts.append("")
        
        # Table of contents
        content_parts.append("## Table of Contents")
        content_parts.append("")
        for chapter in chapters:
            content_parts.append(f"{chapter.order}. [{chapter.title}](#{chapter.title.lower().replace(' ', '-')})")
        content_parts.append("")
        
        # Chapters
        for chapter in chapters:
            if chapter.final_draft_id:
                draft = await self.writer_agent.get_chapter_draft(chapter.final_draft_id)
                if draft:
                    content_parts.append(f"## {chapter.title}")
                    content_parts.append("")
                    content_parts.append(draft.content)
                    content_parts.append("")
        
        # Bibliography
        if include_bibliography:
            content_parts.append("## Bibliography")
            content_parts.append("")
            bibliography = await self._generate_bibliography(book.book_id)
            content_parts.append(bibliography)
        
        # Write to file
        content = "\n".join(content_parts)
        output_path = self.output_directory / f"{book.book_id}.md"
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write(content)
        
        logger.info(f"Exported book to Markdown: {output_path}")
        return str(output_path)
    
    async def _export_docx(
        self,
        book: BookOutline,
        chapters: List[BookChapter],
        include_bibliography: bool
    ) -> str:
        """Export book to DOCX format."""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title page
            title_para = doc.add_heading(book.title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            if book.subtitle:
                subtitle_para = doc.add_heading(book.subtitle, 1)
                subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(f"Author: {book.author}")
            doc.add_paragraph(f"Target Audience: {book.target_audience}")
            doc.add_paragraph("")
            doc.add_paragraph(book.description)
            
            # Add page break
            doc.add_page_break()
            
            # Table of contents
            doc.add_heading("Table of Contents", 1)
            for chapter in chapters:
                doc.add_paragraph(f"{chapter.order}. {chapter.title}")
            
            # Add page break
            doc.add_page_break()
            
            # Chapters
            for chapter in chapters:
                if chapter.final_draft_id:
                    draft = await self.writer_agent.get_chapter_draft(chapter.final_draft_id)
                    if draft:
                        doc.add_heading(chapter.title, 1)
                        
                        # Split content into paragraphs
                        paragraphs = draft.content.split('\n\n')
                        for para_text in paragraphs:
                            if para_text.strip():
                                doc.add_paragraph(para_text.strip())
            
            # Bibliography
            if include_bibliography:
                doc.add_page_break()
                doc.add_heading("Bibliography", 1)
                bibliography = await self._generate_bibliography(book.book_id)
                doc.add_paragraph(bibliography)
            
            # Save document
            output_path = self.output_directory / f"{book.book_id}.docx"
            doc.save(output_path)
            
            logger.info(f"Exported book to DOCX: {output_path}")
            return str(output_path)
            
        except ImportError:
            logger.error("python-docx not available for DOCX export")
            raise ValueError("DOCX export requires python-docx package")
    
    async def _export_pdf(
        self,
        book: BookOutline,
        chapters: List[BookChapter],
        include_bibliography: bool
    ) -> str:
        """Export book to PDF format."""
        try:
            # First export to Markdown, then convert to PDF
            markdown_path = await self._export_markdown(book, chapters, include_bibliography)
            
            # Use pandoc to convert to PDF
            import subprocess
            
            output_path = self.output_directory / f"{book.book_id}.pdf"
            
            subprocess.run([
                "pandoc", markdown_path, "-o", str(output_path),
                "--pdf-engine=pdflatex",
                "--toc", "--toc-depth=2"
            ], check=True)
            
            logger.info(f"Exported book to PDF: {output_path}")
            return str(output_path)
            
        except (ImportError, subprocess.CalledProcessError) as e:
            logger.error(f"PDF export failed: {e}")
            raise ValueError("PDF export requires pandoc and LaTeX")
    
    async def _generate_bibliography(self, book_id: str) -> str:
        """Generate bibliography from used sources."""
        try:
            # Get all chapters for this book
            chapters = [c for c in self.chapters.values() if c.book_id == book_id]
            
            sources = set()
            
            for chapter in chapters:
                if chapter.final_draft_id:
                    draft = await self.writer_agent.get_chapter_draft(chapter.final_draft_id)
                    if draft and draft.research_sources:
                        for source in draft.research_sources:
                            sources.add((source.get("title", ""), source.get("url", ""), source.get("type", "")))
            
            # Format bibliography
            bibliography_items = []
            for i, (title, url, source_type) in enumerate(sorted(sources), 1):
                if source_type == "web" and url:
                    bibliography_items.append(f"{i}. {title}. Available at: {url}")
                else:
                    bibliography_items.append(f"{i}. {title}")
            
            return "\n".join(bibliography_items) if bibliography_items else "No sources cited."
            
        except Exception as e:
            logger.warning(f"Failed to generate bibliography: {e}")
            return "Bibliography generation failed."
    
    async def get_book_status(self, book_id: str) -> Dict[str, Any]:
        """Get book building status."""
        book = self.books.get(book_id)
        if not book:
            return {"error": "Book not found"}
        
        chapters = [c for c in self.chapters.values() if c.book_id == book_id]
        chapters.sort(key=lambda x: x.order)
        
        return {
            "book_id": book_id,
            "title": book.title,
            "total_chapters": len(chapters),
            "completed_chapters": sum(1 for c in chapters if c.status == "final"),
            "in_progress_chapters": sum(1 for c in chapters if c.status == "in_progress"),
            "failed_chapters": sum(1 for c in chapters if c.status == "failed"),
            "chapters": [
                {
                    "chapter_id": c.chapter_id,
                    "title": c.title,
                    "order": c.order,
                    "status": c.status,
                    "word_count": c.word_count,
                    "target_word_count": c.target_word_count
                }
                for c in chapters
            ]
        }
    
    async def get_build_log(self, build_id: str) -> Optional[BuildLog]:
        """Get build log by ID."""
        return self.build_logs.get(build_id)
    
    async def execute_task(self, task_type: str, payload: Dict[str, Any]) -> Any:
        """Execute a book building task."""
        if task_type == "create_book":
            return await self.create_book(
                title=payload.get("title", ""),
                author=payload.get("author", ""),
                description=payload.get("description", ""),
                target_audience=payload.get("target_audience", "general"),
                estimated_word_count=payload.get("estimated_word_count", 50000),
                subtitle=payload.get("subtitle")
            )
        elif task_type == "generate_outline":
            return await self.generate_book_outline(
                book_id=payload.get("book_id"),
                chapter_count=payload.get("chapter_count", 10),
                research_topics=payload.get("research_topics", [])
            )
        elif task_type == "build_book":
            return await self.build_book(
                book_id=payload.get("book_id"),
                chapters_to_build=payload.get("chapters_to_build")
            )
        elif task_type == "export_book":
            return await self.export_book(
                book_id=payload.get("book_id"),
                format=payload.get("format", "markdown"),
                include_bibliography=payload.get("include_bibliography", True)
            )
        elif task_type == "get_status":
            return await self.get_book_status(payload.get("book_id"))
        else:
            raise ValueError(f"Unknown task type: {task_type}")
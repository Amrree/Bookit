"""
Document Ingestor Module

Handles parsing, chunking, metadata extraction, and incremental ingestion of user references.
Supports PDF, MD, TXT, DOCX, EPUB formats.

Chosen libraries:
- PyPDF2: PDF parsing and text extraction
- python-docx: DOCX document processing
- python-epub: EPUB parsing
- pypandoc: Universal document conversion
- pydantic: Data validation and type safety

Adapted from: exp-pj-m-multi-agent-system (https://github.com/krik8235/exp-pj-m-multi-agent-system)
Pattern: Document processing pipeline with metadata extraction
"""

import asyncio
import hashlib
import logging
import mimetypes
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Union
from urllib.parse import urlparse

import pypandoc
import PyPDF2
import pydantic
from docx import Document as DocxDocument
# from epub import Epub  # Not available, using alternative

logger = logging.getLogger(__name__)


class DocumentMetadata(pydantic.BaseModel):
    """Metadata model for ingested documents."""
    source_id: str
    original_filename: str
    file_type: str
    file_size: int
    ingestion_timestamp: datetime
    chunk_count: int
    language: Optional[str] = None
    author: Optional[str] = None
    title: Optional[str] = None
    subject: Optional[str] = None
    keywords: List[str] = []
    checksum: str


class DocumentChunk(pydantic.BaseModel):
    """Model for document chunks."""
    chunk_id: str
    source_id: str
    chunk_index: int
    content: str
    word_count: int
    char_count: int
    metadata: Dict[str, str] = {}


class DocumentIngestor:
    """
    Handles document processing and chunking for RAG ingestion.
    
    Responsibilities:
    - Parse various document formats (PDF, DOCX, TXT, MD, EPUB)
    - Extract metadata and text content
    - Chunk documents for optimal retrieval
    - Generate checksums and provenance data
    - Support incremental ingestion
    """
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        min_chunk_size: int = 100
    ):
        """
        Initialize document ingestor.
        
        Args:
            chunk_size: Target size for document chunks
            chunk_overlap: Overlap between consecutive chunks
            min_chunk_size: Minimum size for a valid chunk
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        
        # Supported file types
        self.supported_types = {
            '.pdf': self._parse_pdf,
            '.docx': self._parse_docx,
            '.txt': self._parse_text,
            '.md': self._parse_markdown,
            '.epub': self._parse_epub
        }
        
        logger.info(f"Document ingestor initialized with chunk_size={chunk_size}")
    
    async def ingest_document(
        self,
        file_path: Union[str, Path],
        source_id: Optional[str] = None
    ) -> Tuple[DocumentMetadata, List[DocumentChunk]]:
        """
        Ingest a document and return metadata and chunks.
        
        Args:
            file_path: Path to document file
            source_id: Optional custom source ID
            
        Returns:
            Tuple of (metadata, chunks)
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        # Generate source ID if not provided
        if not source_id:
            source_id = self._generate_source_id(file_path)
        
        # Get file info
        file_size = file_path.stat().st_size
        file_type = file_path.suffix.lower()
        
        if file_type not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        # Calculate checksum
        checksum = await self._calculate_checksum(file_path)
        
        # Parse document
        content = await self._parse_document(file_path)
        
        # Extract metadata
        metadata = await self._extract_metadata(file_path, content, source_id, file_size, checksum)
        
        # Chunk content
        chunks = await self._chunk_content(content, source_id, metadata)
        
        # Update metadata with chunk count
        metadata.chunk_count = len(chunks)
        
        logger.info(f"Ingested document: {file_path.name} -> {len(chunks)} chunks")
        
        return metadata, chunks
    
    async def _parse_document(self, file_path: Path) -> str:
        """Parse document and extract text content."""
        file_type = file_path.suffix.lower()
        parser = self.supported_types[file_type]
        
        try:
            content = await parser(file_path)
            if not content.strip():
                raise ValueError(f"No content extracted from {file_path}")
            return content
        except Exception as e:
            logger.error(f"Failed to parse {file_path}: {e}")
            raise
    
    async def _parse_pdf(self, file_path: Path) -> str:
        """Parse PDF document."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
                
                return text
        except Exception as e:
            logger.error(f"PDF parsing error: {e}")
            raise
    
    async def _parse_docx(self, file_path: Path) -> str:
        """Parse DOCX document."""
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return text
        except Exception as e:
            logger.error(f"DOCX parsing error: {e}")
            raise
    
    async def _parse_text(self, file_path: Path) -> str:
        """Parse plain text document."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Text parsing error: {e}")
            raise
    
    async def _parse_markdown(self, file_path: Path) -> str:
        """Parse Markdown document."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # Convert markdown to plain text using pypandoc
            try:
                text = pypandoc.convert_text(content, 'plain', format='md')
                return text
            except Exception:
                # Fallback to raw content if pypandoc fails
                return content
        except Exception as e:
            logger.error(f"Markdown parsing error: {e}")
            raise
    
    async def _parse_epub(self, file_path: Path) -> str:
        """Parse EPUB document."""
        try:
            # This is a simplified implementation
            # In practice, you would use a proper EPUB library
            with open(file_path, 'rb') as file:
                # Basic text extraction - this is not a complete EPUB parser
                content = file.read().decode('utf-8', errors='ignore')
                return content
        except Exception as e:
            logger.error(f"EPUB parsing error: {e}")
            raise
    
    async def _extract_metadata(
        self,
        file_path: Path,
        content: str,
        source_id: str,
        file_size: int,
        checksum: str
    ) -> DocumentMetadata:
        """Extract metadata from document."""
        # Basic metadata extraction
        metadata = DocumentMetadata(
            source_id=source_id,
            original_filename=file_path.name,
            file_type=file_path.suffix.lower(),
            file_size=file_size,
            ingestion_timestamp=datetime.now(),
            chunk_count=0,  # Will be updated after chunking
            checksum=checksum
        )
        
        # Try to extract additional metadata based on file type
        if file_path.suffix.lower() == '.pdf':
            metadata = await self._extract_pdf_metadata(file_path, metadata)
        elif file_path.suffix.lower() == '.docx':
            metadata = await self._extract_docx_metadata(file_path, metadata)
        
        return metadata
    
    async def _extract_pdf_metadata(self, file_path: Path, metadata: DocumentMetadata) -> DocumentMetadata:
        """Extract PDF-specific metadata."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                info = pdf_reader.metadata
                
                if info:
                    metadata.title = info.get('/Title', '')
                    metadata.author = info.get('/Author', '')
                    metadata.subject = info.get('/Subject', '')
                    metadata.keywords = info.get('/Keywords', '').split(',') if info.get('/Keywords') else []
        except Exception as e:
            logger.warning(f"Failed to extract PDF metadata: {e}")
        
        return metadata
    
    async def _extract_docx_metadata(self, file_path: Path, metadata: DocumentMetadata) -> DocumentMetadata:
        """Extract DOCX-specific metadata."""
        try:
            doc = DocxDocument(file_path)
            core_props = doc.core_properties
            
            metadata.title = core_props.title or ''
            metadata.author = core_props.author or ''
            metadata.subject = core_props.subject or ''
            metadata.keywords = core_props.keywords.split(',') if core_props.keywords else []
        except Exception as e:
            logger.warning(f"Failed to extract DOCX metadata: {e}")
        
        return metadata
    
    async def _chunk_content(
        self,
        content: str,
        source_id: str,
        metadata: DocumentMetadata
    ) -> List[DocumentChunk]:
        """Chunk content into smaller pieces."""
        if not content.strip():
            return []
        
        # Split content into sentences for better chunking
        sentences = self._split_into_sentences(content)
        
        chunks = []
        current_chunk = ""
        chunk_index = 0
        
        for sentence in sentences:
            # Check if adding this sentence would exceed chunk size
            if len(current_chunk) + len(sentence) > self.chunk_size and current_chunk:
                # Create chunk if it meets minimum size
                if len(current_chunk) >= self.min_chunk_size:
                    chunk = self._create_chunk(
                        current_chunk, source_id, chunk_index, metadata
                    )
                    chunks.append(chunk)
                    chunk_index += 1
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + sentence
            else:
                current_chunk += sentence
        
        # Add final chunk if it meets minimum size
        if len(current_chunk) >= self.min_chunk_size:
            chunk = self._create_chunk(current_chunk, source_id, chunk_index, metadata)
            chunks.append(chunk)
        
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences."""
        import re
        
        # Simple sentence splitting
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        # Add sentence endings back
        sentences_with_endings = []
        for i, sentence in enumerate(sentences):
            if i < len(sentences) - 1:
                sentences_with_endings.append(sentence + '.')
            else:
                sentences_with_endings.append(sentence)
        
        return sentences_with_endings
    
    def _get_overlap_text(self, text: str) -> str:
        """Get overlap text from the end of current chunk."""
        if len(text) <= self.chunk_overlap:
            return text
        
        # Get last chunk_overlap characters
        overlap = text[-self.chunk_overlap:]
        
        # Try to break at word boundary
        words = overlap.split()
        if len(words) > 1:
            return ' '.join(words[1:]) + ' '
        
        return overlap
    
    def _create_chunk(
        self,
        content: str,
        source_id: str,
        chunk_index: int,
        metadata: DocumentMetadata
    ) -> DocumentChunk:
        """Create a document chunk."""
        chunk_id = f"{source_id}_chunk_{chunk_index}"
        
        return DocumentChunk(
            chunk_id=chunk_id,
            source_id=source_id,
            chunk_index=chunk_index,
            content=content.strip(),
            word_count=len(content.split()),
            char_count=len(content),
            metadata={
                "original_filename": metadata.original_filename,
                "file_type": metadata.file_type,
                "chunk_size": str(len(content))
            }
        )
    
    async def _calculate_checksum(self, file_path: Path) -> str:
        """Calculate file checksum."""
        hash_md5 = hashlib.md5()
        
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        
        return hash_md5.hexdigest()
    
    def _generate_source_id(self, file_path: Path) -> str:
        """Generate unique source ID for document."""
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"{file_path.stem}_{timestamp}"
    
    def get_supported_types(self) -> List[str]:
        """Get list of supported file types."""
        return list(self.supported_types.keys())
    
    def is_supported(self, file_path: Union[str, Path]) -> bool:
        """Check if file type is supported."""
        file_path = Path(file_path)
        return file_path.suffix.lower() in self.supported_types
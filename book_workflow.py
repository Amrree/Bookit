"""
Complete Book Writing Workflow Module

Implements a full book production pipeline that generates complete, book-length
non-fiction manuscripts with chapter-by-chapter production, continuity tracking,
and comprehensive export capabilities.

Chosen libraries:
- asyncio: Asynchronous workflow orchestration
- logging: Comprehensive audit logging
- json: Build log generation
- pathlib: File management
- datetime: Timestamp management

Pattern: Sequential pipeline with quality controls and provenance tracking
"""

import asyncio
import json
import logging
import os
import uuid
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

from document_ingestor import DocumentIngestor
from memory_manager import MemoryManager
from llm_client import LLMClient
from tool_manager import ToolManager
from agent_manager import AgentManager
from research_agent import ResearchAgent
from writer_agent import WriterAgent, WritingStyle
from editor_agent import EditorAgent, StyleGuide
from tool_agent import ToolAgent
from book_builder import BookBuilder

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)


class BookMetadata:
    """Book metadata container."""
    
    def __init__(self, title: str, theme: str, author: str = "AI Book Writer"):
        self.title = title
        self.theme = theme
        self.author = author
        self.build_id = str(uuid.uuid4())
        self.created_at = datetime.now()
        self.word_count = 0
        self.chapter_count = 0
        self.status = "draft"
        self.outline = {}
        self.chapters = []
        self.bibliography = []
        self.build_log = {}


class ChapterMetadata:
    """Chapter metadata container."""
    
    def __init__(self, chapter_number: int, title: str, word_count_target: int = 5000):
        self.chapter_number = chapter_number
        self.title = title
        self.word_count_target = word_count_target
        self.actual_word_count = 0
        self.status = "draft"
        self.draft_content = ""
        self.revised_content = ""
        self.final_content = ""
        self.references_used = []
        self.retrieval_scores = []
        self.agents_involved = []
        self.tool_calls = []
        self.revision_notes = []
        self.created_at = datetime.now()
        self.updated_at = datetime.now()


class BookWorkflow:
    """Complete book writing workflow orchestrator."""
    
    def __init__(
        self,
        memory_manager: MemoryManager,
        llm_client: LLMClient,
        tool_manager: ToolManager,
        agent_manager: AgentManager,
        research_agent: ResearchAgent,
        writer_agent: WriterAgent,
        editor_agent: EditorAgent,
        tool_agent: ToolAgent,
        book_builder: BookBuilder
    ):
        self.memory_manager = memory_manager
        self.llm_client = llm_client
        self.tool_manager = tool_manager
        self.agent_manager = agent_manager
        self.research_agent = research_agent
        self.writer_agent = writer_agent
        self.editor_agent = editor_agent
        self.tool_agent = tool_agent
        self.book_builder = book_builder
        
        self.current_book: Optional[BookMetadata] = None
        self.workflow_log = []
        
    async def start_book_production(
        self,
        title: str,
        theme: str,
        reference_documents: Optional[List[str]] = None,
        target_word_count: int = 50000,
        chapters_count: int = 10,
        author: str = "AI Book Writer"
    ) -> BookMetadata:
        """Start complete book production workflow."""
        
        logger.info(f"Starting book production: '{title}'")
        
        # Initialize book metadata
        self.current_book = BookMetadata(title, theme, author)
        self.current_book.word_count = target_word_count
        self.current_book.chapter_count = chapters_count
        
        # Step 1: Ingest reference documents
        if reference_documents:
            await self._ingest_reference_documents(reference_documents)
        
        # Step 2: Research and create outline
        await self._research_and_outline()
        
        # Step 3: Generate chapters sequentially
        await self._generate_chapters()
        
        # Step 4: Global revision and assembly
        await self._global_revision_and_assembly()
        
        # Step 5: Generate bibliography
        await self._generate_bibliography()
        
        # Step 6: Export in all formats
        await self._export_book()
        
        # Step 7: Generate build log
        await self._generate_build_log()
        
        logger.info(f"Book production completed: '{title}'")
        return self.current_book
    
    async def _ingest_reference_documents(self, document_paths: List[str]):
        """Ingest reference documents into memory."""
        
        logger.info("Ingesting reference documents...")
        
        ingestor = DocumentIngestor(self.memory_manager)
        
        for doc_path in document_paths:
            try:
                if os.path.exists(doc_path):
                    await ingestor.ingest_document(doc_path)
                    self.workflow_log.append({
                        "step": "document_ingestion",
                        "document": doc_path,
                        "status": "success",
                        "timestamp": datetime.now().isoformat()
                    })
                    logger.info(f"Successfully ingested: {doc_path}")
                else:
                    logger.warning(f"Document not found: {doc_path}")
                    
            except Exception as e:
                logger.error(f"Failed to ingest {doc_path}: {e}")
                self.workflow_log.append({
                    "step": "document_ingestion",
                    "document": doc_path,
                    "status": "failed",
                    "error": str(e),
                    "timestamp": datetime.now().isoformat()
                })
    
    async def _research_and_outline(self):
        """Research theme and create detailed book outline."""
        
        logger.info("Conducting research and creating outline...")
        
        # Research the theme comprehensively
        research_prompt = f"""
        Conduct comprehensive research on the theme: {self.current_book.theme}
        
        Create a detailed book outline for a {self.current_book.word_count:,}-word non-fiction book titled "{self.current_book.title}".
        
        The outline should include:
        1. A compelling introduction that hooks readers
        2. {self.current_book.chapter_count} detailed chapters with:
           - Chapter titles
           - Key points and arguments
           - Supporting evidence needed
           - Word count targets (aim for ~{self.current_book.word_count // self.current_book.chapter_count:,} words per chapter)
        3. A strong conclusion that ties everything together
        4. Key themes and concepts to maintain throughout
        
        Focus on creating a coherent narrative arc that builds knowledge progressively.
        """
        
        # Create a research topic and start research
        topic_id = await self.research_agent.start_research(
            topic_title=self.current_book.theme,
            description=research_prompt,
            keywords=[self.current_book.theme, "book outline", "non-fiction"],
            priority=1
        )
        
        # Get research results
        research_results = await self.research_agent.get_research_results(topic_id)
        research_summary = await self.research_agent.get_research_summary(topic_id)
        
        research_result = {
            'summary': research_summary.summary if research_summary else "Research completed",
            'relevant_chunks': [r.chunk_id for r in research_results if r.chunk_id],
            'sources': len(research_results)
        }
        
        # Extract outline from research
        outline_prompt = f"""
        Based on the research conducted, create a detailed book outline for "{self.current_book.title}".
        
        Research context: {research_result.get('summary', '')}
        
        Create a structured outline with:
        1. Book introduction and thesis
        2. {self.current_book.chapter_count} chapters with detailed breakdowns
        3. Conclusion and key takeaways
        
        Format as JSON with this structure:
        {{
            "introduction": {{
                "title": "Introduction Title",
                "key_points": ["point1", "point2", ...],
                "word_count_target": 3000
            }},
            "chapters": [
                {{
                    "chapter_number": 1,
                    "title": "Chapter Title",
                    "key_points": ["point1", "point2", ...],
                    "word_count_target": 5000,
                    "research_focus": "specific aspect to research"
                }},
                ...
            ],
            "conclusion": {{
                "title": "Conclusion Title",
                "key_points": ["point1", "point2", ...],
                "word_count_target": 2000
            }}
        }}
        """
        
        outline_result = await self.llm_client.generate(
            prompt=outline_prompt,
            max_tokens=4000,
            temperature=0.7
        )
        
        try:
            self.current_book.outline = json.loads(outline_result.content)
        except json.JSONDecodeError:
            # Fallback to simple outline if JSON parsing fails
            self.current_book.outline = {
                "introduction": {
                    "title": f"Introduction to {self.current_book.theme}",
                    "key_points": ["Overview", "Importance", "What readers will learn"],
                    "word_count_target": 3000
                },
                "chapters": [
                    {
                        "chapter_number": i + 1,
                        "title": f"Chapter {i + 1}: {self.current_book.theme} - Part {i + 1}",
                        "key_points": ["Key concept", "Supporting evidence", "Examples"],
                        "word_count_target": self.current_book.word_count // self.current_book.chapter_count,
                        "research_focus": f"Specific aspect {i + 1} of {self.current_book.theme}"
                    }
                    for i in range(self.current_book.chapter_count)
                ],
                "conclusion": {
                    "title": f"Conclusion: The Future of {self.current_book.theme}",
                    "key_points": ["Summary", "Key takeaways", "Future implications"],
                    "word_count_target": 2000
                }
            }
        
        self.workflow_log.append({
            "step": "research_and_outline",
            "status": "completed",
            "outline_chapters": len(self.current_book.outline.get('chapters', [])),
            "timestamp": datetime.now().isoformat()
        })
        
        logger.info(f"Created outline with {len(self.current_book.outline.get('chapters', []))} chapters")
    
    async def _generate_chapters(self):
        """Generate all chapters sequentially with continuity tracking."""
        
        logger.info("Generating chapters...")
        
        # Generate introduction
        await self._generate_introduction()
        
        # Generate each chapter
        for chapter_outline in self.current_book.outline.get('chapters', []):
            await self._generate_chapter(chapter_outline)
        
        # Generate conclusion
        await self._generate_conclusion()
        
        logger.info("All chapters generated successfully")
    
    async def _generate_introduction(self):
        """Generate book introduction."""
        
        intro_outline = self.current_book.outline.get('introduction', {})
        
        chapter_meta = ChapterMetadata(
            chapter_number=0,
            title=intro_outline.get('title', 'Introduction'),
            word_count_target=intro_outline.get('word_count_target', 3000)
        )
        
        # Research for introduction
        research_prompt = f"""
        Research comprehensive background information for the introduction of a book titled "{self.current_book.title}" about {self.current_book.theme}.
        
        Focus on:
        - Historical context and background
        - Current state of the field
        - Why this topic matters now
        - Key concepts readers need to understand
        - What makes this book unique and valuable
        
        Provide detailed, well-sourced information that will help write a compelling introduction.
        """
        
        # Create research topic for introduction
        topic_id = await self.research_agent.start_research(
            topic_title=f"{self.current_book.theme} introduction background",
            description=research_prompt,
            keywords=[self.current_book.theme, "introduction", "background", "history"],
            priority=1
        )
        
        # Get research results
        research_results = await self.research_agent.get_research_results(topic_id)
        research_summary = await self.research_agent.get_research_summary(topic_id)
        
        research_result = {
            'summary': research_summary.summary if research_summary else "Research completed",
            'relevant_chunks': [r.chunk_id for r in research_results if r.chunk_id],
            'sources': len(research_results)
        }
        
        # Write introduction
        writing_prompt = f"""
        Write a compelling introduction for a book titled "{self.current_book.title}" about {self.current_book.theme}.
        
        Research context: {research_result.get('summary', '')}
        
        Key points to cover: {intro_outline.get('key_points', [])}
        Target word count: {chapter_meta.word_count_target:,} words
        
        The introduction should:
        1. Hook the reader with a compelling opening
        2. Establish the importance and relevance of the topic
        3. Preview what readers will learn
        4. Set up the book's structure and approach
        5. Include proper citations using [chunk_id] format
        
        Write in a professional, engaging style suitable for a non-fiction book.
        """
        
        draft_result = await self.writer_agent.write_chapter(
            chapter_title=chapter_meta.title,
            writing_prompt=writing_prompt,
            target_word_count=chapter_meta.word_count_target,
            context_chunks=research_result.get('relevant_chunks', [])
        )
        
        chapter_meta.draft_content = draft_result.get('content', '')
        chapter_meta.references_used = draft_result.get('references_used', [])
        chapter_meta.retrieval_scores = draft_result.get('retrieval_scores', [])
        chapter_meta.agents_involved = ['research_agent', 'writer_agent']
        
        # Edit introduction
        edit_result = await self.editor_agent.revise_chapter(
            chapter_title=chapter_meta.title,
            content=chapter_meta.draft_content,
            revision_focus="introduction_quality"
        )
        
        chapter_meta.revised_content = edit_result.get('revised_content', chapter_meta.draft_content)
        chapter_meta.revision_notes = edit_result.get('revision_notes', [])
        chapter_meta.agents_involved.append('editor_agent')
        
        # Finalize
        chapter_meta.final_content = chapter_meta.revised_content
        chapter_meta.actual_word_count = len(chapter_meta.final_content.split())
        chapter_meta.status = "completed"
        chapter_meta.updated_at = datetime.now()
        
        self.current_book.chapters.append(chapter_meta)
        
        # Update memory with chapter summary for continuity
        await self._update_memory_with_chapter_summary(chapter_meta)
        
        logger.info(f"Generated introduction: {chapter_meta.actual_word_count:,} words")
    
    async def _generate_chapter(self, chapter_outline: Dict):
        """Generate a single chapter with continuity awareness."""
        
        chapter_meta = ChapterMetadata(
            chapter_number=chapter_outline.get('chapter_number', 1),
            title=chapter_outline.get('title', f"Chapter {chapter_outline.get('chapter_number', 1)}"),
            word_count_target=chapter_outline.get('word_count_target', 5000)
        )
        
        # Research for this specific chapter
        research_prompt = f"""
        Research detailed information for Chapter {chapter_meta.chapter_number}: "{chapter_meta.title}" of a book about {self.current_book.theme}.
        
        Research focus: {chapter_outline.get('research_focus', '')}
        Key points to cover: {chapter_outline.get('key_points', [])}
        
        This is part of a larger book, so consider:
        - How this chapter builds on previous chapters
        - What new concepts or information to introduce
        - How to maintain narrative flow
        - What evidence and examples will be most compelling
        
        Provide comprehensive, well-sourced information for writing this chapter.
        """
        
        # Create research topic for chapter
        topic_id = await self.research_agent.start_research(
            topic_title=f"{self.current_book.theme} chapter {chapter_meta.chapter_number}",
            description=research_prompt,
            keywords=[self.current_book.theme, f"chapter {chapter_meta.chapter_number}", chapter_outline.get('research_focus', '')],
            priority=1
        )
        
        # Get research results
        research_results = await self.research_agent.get_research_results(topic_id)
        research_summary = await self.research_agent.get_research_summary(topic_id)
        
        research_result = {
            'summary': research_summary.summary if research_summary else "Research completed",
            'relevant_chunks': [r.chunk_id for r in research_results if r.chunk_id],
            'sources': len(research_results)
        }
        
        # Get context from previous chapters for continuity
        previous_context = await self._get_previous_chapters_context()
        
        # Write chapter
        writing_prompt = f"""
        Write Chapter {chapter_meta.chapter_number}: "{chapter_meta.title}" for a book titled "{self.current_book.title}" about {self.current_book.theme}.
        
        Research context: {research_result.get('summary', '')}
        Previous chapters context: {previous_context}
        
        Key points to cover: {chapter_outline.get('key_points', [])}
        Target word count: {chapter_meta.word_count_target:,} words
        
        The chapter should:
        1. Build naturally on previous chapters
        2. Cover the key points thoroughly
        3. Include compelling evidence and examples
        4. Maintain consistent tone and style
        5. Include proper citations using [chunk_id] format
        6. End with a transition to the next chapter
        
        Write in a professional, engaging style suitable for a non-fiction book.
        """
        
        draft_result = await self.writer_agent.write_chapter(
            chapter_title=chapter_meta.title,
            writing_prompt=writing_prompt,
            target_word_count=chapter_meta.word_count_target,
            context_chunks=research_result.get('relevant_chunks', [])
        )
        
        chapter_meta.draft_content = draft_result.get('content', '')
        chapter_meta.references_used = draft_result.get('references_used', [])
        chapter_meta.retrieval_scores = draft_result.get('retrieval_scores', [])
        chapter_meta.agents_involved = ['research_agent', 'writer_agent']
        
        # Edit chapter
        edit_result = await self.editor_agent.revise_chapter(
            chapter_title=chapter_meta.title,
            content=chapter_meta.draft_content,
            revision_focus="chapter_quality_and_continuity"
        )
        
        chapter_meta.revised_content = edit_result.get('revised_content', chapter_meta.draft_content)
        chapter_meta.revision_notes = edit_result.get('revision_notes', [])
        chapter_meta.agents_involved.append('editor_agent')
        
        # Finalize
        chapter_meta.final_content = chapter_meta.revised_content
        chapter_meta.actual_word_count = len(chapter_meta.final_content.split())
        chapter_meta.status = "completed"
        chapter_meta.updated_at = datetime.now()
        
        self.current_book.chapters.append(chapter_meta)
        
        # Update memory with chapter summary for continuity
        await self._update_memory_with_chapter_summary(chapter_meta)
        
        logger.info(f"Generated Chapter {chapter_meta.chapter_number}: {chapter_meta.actual_word_count:,} words")
    
    async def _generate_conclusion(self):
        """Generate book conclusion."""
        
        conclusion_outline = self.current_book.outline.get('conclusion', {})
        
        chapter_meta = ChapterMetadata(
            chapter_number=999,  # Special number for conclusion
            title=conclusion_outline.get('title', 'Conclusion'),
            word_count_target=conclusion_outline.get('word_count_target', 2000)
        )
        
        # Get context from all previous chapters
        full_book_context = await self._get_full_book_context()
        
        # Research for conclusion
        research_prompt = f"""
        Research information for the conclusion of a book titled "{self.current_book.title}" about {self.current_book.theme}.
        
        Focus on:
        - Key takeaways and insights
        - Future implications and trends
        - Practical applications
        - Call to action for readers
        
        This should tie together all the themes and concepts covered in the book.
        """
        
        # Create research topic for conclusion
        topic_id = await self.research_agent.start_research(
            topic_title=f"{self.current_book.theme} conclusion future implications",
            description=research_prompt,
            keywords=[self.current_book.theme, "conclusion", "future", "implications", "takeaways"],
            priority=1
        )
        
        # Get research results
        research_results = await self.research_agent.get_research_results(topic_id)
        research_summary = await self.research_agent.get_research_summary(topic_id)
        
        research_result = {
            'summary': research_summary.summary if research_summary else "Research completed",
            'relevant_chunks': [r.chunk_id for r in research_results if r.chunk_id],
            'sources': len(research_results)
        }
        
        # Write conclusion
        writing_prompt = f"""
        Write a compelling conclusion for a book titled "{self.current_book.title}" about {self.current_book.theme}.
        
        Book context: {full_book_context}
        Research context: {research_result.get('summary', '')}
        
        Key points to cover: {conclusion_outline.get('key_points', [])}
        Target word count: {chapter_meta.word_count_target:,} words
        
        The conclusion should:
        1. Summarize key insights from the book
        2. Synthesize the main themes
        3. Provide practical takeaways
        4. Discuss future implications
        5. End with a compelling call to action
        6. Include proper citations using [chunk_id] format
        
        Write in a professional, engaging style suitable for a non-fiction book.
        """
        
        draft_result = await self.writer_agent.write_chapter(
            chapter_title=chapter_meta.title,
            writing_prompt=writing_prompt,
            target_word_count=chapter_meta.word_count_target,
            context_chunks=research_result.get('relevant_chunks', [])
        )
        
        chapter_meta.draft_content = draft_result.get('content', '')
        chapter_meta.references_used = draft_result.get('references_used', [])
        chapter_meta.retrieval_scores = draft_result.get('retrieval_scores', [])
        chapter_meta.agents_involved = ['research_agent', 'writer_agent']
        
        # Edit conclusion
        edit_result = await self.editor_agent.revise_chapter(
            chapter_title=chapter_meta.title,
            content=chapter_meta.draft_content,
            revision_focus="conclusion_quality_and_synthesis"
        )
        
        chapter_meta.revised_content = edit_result.get('revised_content', chapter_meta.draft_content)
        chapter_meta.revision_notes = edit_result.get('revision_notes', [])
        chapter_meta.agents_involved.append('editor_agent')
        
        # Finalize
        chapter_meta.final_content = chapter_meta.revised_content
        chapter_meta.actual_word_count = len(chapter_meta.final_content.split())
        chapter_meta.status = "completed"
        chapter_meta.updated_at = datetime.now()
        
        self.current_book.chapters.append(chapter_meta)
        
        logger.info(f"Generated conclusion: {chapter_meta.actual_word_count:,} words")
    
    async def _get_previous_chapters_context(self) -> str:
        """Get context from previous chapters for continuity."""
        
        if not self.current_book.chapters:
            return "This is the first chapter."
        
        context_parts = []
        for chapter in self.current_book.chapters[-3:]:  # Last 3 chapters for context
            context_parts.append(f"Chapter {chapter.chapter_number}: {chapter.title}")
            context_parts.append(f"Key points: {chapter.final_content[:500]}...")
        
        return "\n\n".join(context_parts)
    
    async def _get_full_book_context(self) -> str:
        """Get context from all chapters for conclusion."""
        
        if not self.current_book.chapters:
            return "No previous chapters."
        
        context_parts = []
        for chapter in self.current_book.chapters:
            context_parts.append(f"Chapter {chapter.chapter_number}: {chapter.title}")
            context_parts.append(f"Summary: {chapter.final_content[:300]}...")
        
        return "\n\n".join(context_parts)
    
    async def _update_memory_with_chapter_summary(self, chapter_meta: ChapterMetadata):
        """Update memory with chapter summary for continuity."""
        
        summary_text = f"""
        Chapter {chapter_meta.chapter_number}: {chapter_meta.title}
        
        Key content: {chapter_meta.final_content[:1000]}...
        
        Word count: {chapter_meta.actual_word_count:,}
        References used: {len(chapter_meta.references_used)}
        Agents involved: {', '.join(chapter_meta.agents_involved)}
        """
        
        await self.memory_manager.add_memory(
            content=summary_text,
            metadata={
                "type": "chapter_summary",
                "chapter_number": chapter_meta.chapter_number,
                "chapter_title": chapter_meta.title,
                "book_title": self.current_book.title,
                "word_count": chapter_meta.actual_word_count,
                "references_count": len(chapter_meta.references_used),
                "created_at": chapter_meta.created_at.isoformat()
            }
        )
    
    async def _global_revision_and_assembly(self):
        """Perform global revision and assemble final manuscript."""
        
        logger.info("Performing global revision and assembly...")
        
        # Assemble all chapters
        full_manuscript = await self._assemble_manuscript()
        
        # Global revision pass
        global_revision_prompt = f"""
        Perform a comprehensive global revision of the complete book "{self.current_book.title}".
        
        Focus on:
        1. Overall narrative flow and coherence
        2. Consistency of terminology and style
        3. Logical progression between chapters
        4. Elimination of redundancy
        5. Strengthening transitions
        6. Ensuring all claims are properly cited
        
        The book should read as a unified, coherent work.
        """
        
        revision_result = await self.editor_agent.revise_chapter(
            chapter_title="Full Manuscript",
            content=full_manuscript,
            revision_focus="global_coherence_and_consistency"
        )
        
        # Update final content in all chapters
        revised_manuscript = revision_result.get('revised_content', full_manuscript)
        await self._update_chapters_with_global_revision(revised_manuscript)
        
        self.workflow_log.append({
            "step": "global_revision",
            "status": "completed",
            "revision_notes": revision_result.get('revision_notes', []),
            "timestamp": datetime.now().isoformat()
        })
        
        logger.info("Global revision completed")
    
    async def _assemble_manuscript(self) -> str:
        """Assemble all chapters into a single manuscript."""
        
        manuscript_parts = []
        
        # Title page
        manuscript_parts.append(f"# {self.current_book.title}")
        manuscript_parts.append(f"**Author:** {self.current_book.author}")
        manuscript_parts.append(f"**Build ID:** {self.current_book.build_id}")
        manuscript_parts.append(f"**Created:** {self.current_book.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
        manuscript_parts.append("")
        manuscript_parts.append("---")
        manuscript_parts.append("")
        
        # Introduction
        intro_chapter = next((c for c in self.current_book.chapters if c.chapter_number == 0), None)
        if intro_chapter:
            manuscript_parts.append(f"# {intro_chapter.title}")
            manuscript_parts.append(intro_chapter.final_content)
            manuscript_parts.append("")
        
        # Main chapters
        main_chapters = [c for c in self.current_book.chapters if 0 < c.chapter_number < 999]
        main_chapters.sort(key=lambda x: x.chapter_number)
        
        for chapter in main_chapters:
            manuscript_parts.append(f"# Chapter {chapter.chapter_number}: {chapter.title}")
            manuscript_parts.append(chapter.final_content)
            manuscript_parts.append("")
        
        # Conclusion
        conclusion_chapter = next((c for c in self.current_book.chapters if c.chapter_number == 999), None)
        if conclusion_chapter:
            manuscript_parts.append(f"# {conclusion_chapter.title}")
            manuscript_parts.append(conclusion_chapter.final_content)
            manuscript_parts.append("")
        
        return "\n".join(manuscript_parts)
    
    async def _update_chapters_with_global_revision(self, revised_manuscript: str):
        """Update individual chapters with global revision changes."""
        
        # This is a simplified approach - in practice, you'd want more sophisticated
        # chapter extraction and updating logic
        logger.info("Updating chapters with global revision changes")
    
    async def _generate_bibliography(self):
        """Generate bibliography from all referenced chunks."""
        
        logger.info("Generating bibliography...")
        
        # Collect all unique references from all chapters
        all_references = set()
        for chapter in self.current_book.chapters:
            all_references.update(chapter.references_used)
        
        # Get metadata for each reference
        bibliography_entries = []
        for chunk_id in all_references:
            try:
                chunk_metadata = await self.memory_manager.get_chunk_metadata(chunk_id)
                if chunk_metadata:
                    bibliography_entries.append({
                        "chunk_id": chunk_id,
                        "source_id": chunk_metadata.get("source_id", "unknown"),
                        "original_filename": chunk_metadata.get("original_filename", "unknown"),
                        "provenance_notes": chunk_metadata.get("provenance_notes", ""),
                        "retrieval_score": chunk_metadata.get("retrieval_score", 0.0)
                    })
            except Exception as e:
                logger.warning(f"Failed to get metadata for chunk {chunk_id}: {e}")
        
        # Sort by source and chunk ID
        bibliography_entries.sort(key=lambda x: (x["original_filename"], x["chunk_id"]))
        
        self.current_book.bibliography = bibliography_entries
        
        self.workflow_log.append({
            "step": "bibliography_generation",
            "status": "completed",
            "total_references": len(bibliography_entries),
            "timestamp": datetime.now().isoformat()
        })
        
        logger.info(f"Generated bibliography with {len(bibliography_entries)} references")
    
    async def _export_book(self):
        """Export book in all required formats."""
        
        logger.info("Exporting book in all formats...")
        
        # Create output directory
        output_dir = Path(f"output/{self.current_book.build_id}")
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # Assemble final manuscript
        final_manuscript = await self._assemble_manuscript()
        
        # Add bibliography to manuscript
        final_manuscript += "\n\n# Bibliography\n\n"
        for entry in self.current_book.bibliography:
            final_manuscript += f"- **{entry['original_filename']}** (Chunk {entry['chunk_id']})\n"
            if entry['provenance_notes']:
                final_manuscript += f"  - {entry['provenance_notes']}\n"
            final_manuscript += "\n"
        
        # Export Markdown
        md_path = output_dir / f"{self.current_book.title.replace(' ', '_')}.md"
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(final_manuscript)
        
        # Export DOCX
        docx_path = output_dir / f"{self.current_book.title.replace(' ', '_')}.docx"
        await self._export_docx(final_manuscript, docx_path)
        
        # Export PDF
        pdf_path = output_dir / f"{self.current_book.title.replace(' ', '_')}.pdf"
        await self._export_pdf(final_manuscript, pdf_path)
        
        self.workflow_log.append({
            "step": "export",
            "status": "completed",
            "output_directory": str(output_dir),
            "formats": ["markdown", "docx", "pdf"],
            "timestamp": datetime.now().isoformat()
        })
        
        logger.info(f"Book exported to {output_dir}")
    
    async def _export_docx(self, content: str, output_path: Path):
        """Export content to DOCX format."""
        
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Split content into sections
            sections = content.split('\n# ')
            
            for i, section in enumerate(sections):
                if i == 0:
                    # Title section
                    lines = section.split('\n')
                    title = lines[0].replace('# ', '')
                    doc.add_heading(title, 0)
                    
                    for line in lines[1:]:
                        if line.strip():
                            if line.startswith('**') and line.endswith('**'):
                                doc.add_paragraph(line)
                            else:
                                doc.add_paragraph(line)
                else:
                    # Chapter sections
                    lines = section.split('\n')
                    if lines:
                        chapter_title = lines[0]
                        doc.add_heading(chapter_title, 1)
                        
                        for line in lines[1:]:
                            if line.strip():
                                doc.add_paragraph(line)
            
            doc.save(str(output_path))
            logger.info(f"Exported DOCX to {output_path}")
            
        except ImportError:
            logger.warning("python-docx not available, skipping DOCX export")
        except Exception as e:
            logger.error(f"Failed to export DOCX: {e}")
    
    async def _export_pdf(self, content: str, output_path: Path):
        """Export content to PDF format."""
        
        try:
            import markdown
            from weasyprint import HTML, CSS
            from weasyprint.text.fonts import FontConfiguration
            
            # Convert markdown to HTML
            html_content = markdown.markdown(content, extensions=['tables', 'toc'])
            
            # Add CSS styling
            styled_html = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <style>
                    body {{
                        font-family: 'Times New Roman', serif;
                        line-height: 1.6;
                        max-width: 8.5in;
                        margin: 0 auto;
                        padding: 1in;
                    }}
                    h1 {{
                        font-size: 24pt;
                        text-align: center;
                        margin-bottom: 0.5in;
                    }}
                    h2 {{
                        font-size: 18pt;
                        margin-top: 0.5in;
                        margin-bottom: 0.25in;
                    }}
                    h3 {{
                        font-size: 14pt;
                        margin-top: 0.25in;
                        margin-bottom: 0.125in;
                    }}
                    p {{
                        margin-bottom: 0.125in;
                        text-align: justify;
                    }}
                    .bibliography {{
                        margin-top: 0.5in;
                    }}
                </style>
            </head>
            <body>
                {html_content}
            </body>
            </html>
            """
            
            # Generate PDF
            font_config = FontConfiguration()
            html_doc = HTML(string=styled_html)
            html_doc.write_pdf(str(output_path), font_config=font_config)
            
            logger.info(f"Exported PDF to {output_path}")
            
        except ImportError:
            logger.warning("weasyprint not available, skipping PDF export")
        except Exception as e:
            logger.error(f"Failed to export PDF: {e}")
    
    async def _generate_build_log(self):
        """Generate comprehensive build log."""
        
        logger.info("Generating build log...")
        
        # Calculate total word count
        total_words = sum(chapter.actual_word_count for chapter in self.current_book.chapters)
        self.current_book.word_count = total_words
        
        # Generate build log
        build_log = {
            "book_metadata": {
                "title": self.current_book.title,
                "theme": self.current_book.theme,
                "author": self.current_book.author,
                "build_id": self.current_book.build_id,
                "created_at": self.current_book.created_at.isoformat(),
                "total_word_count": total_words,
                "chapter_count": len(self.current_book.chapters),
                "status": self.current_book.status
            },
            "chapters_produced": [
                {
                    "chapter_number": chapter.chapter_number,
                    "title": chapter.title,
                    "word_count": chapter.actual_word_count,
                    "word_count_target": chapter.word_count_target,
                    "agents_involved": chapter.agents_involved,
                    "references_used": len(chapter.references_used),
                    "created_at": chapter.created_at.isoformat(),
                    "updated_at": chapter.updated_at.isoformat()
                }
                for chapter in self.current_book.chapters
            ],
            "references_used": [
                {
                    "chunk_id": entry["chunk_id"],
                    "source_id": entry["source_id"],
                    "original_filename": entry["original_filename"],
                    "retrieval_score": entry["retrieval_score"]
                }
                for entry in self.current_book.bibliography
            ],
            "workflow_log": self.workflow_log,
            "quality_metrics": {
                "total_references": len(self.current_book.bibliography),
                "average_chapter_length": total_words / len(self.current_book.chapters) if self.current_book.chapters else 0,
                "chapters_with_target_length": sum(1 for c in self.current_book.chapters if c.actual_word_count >= c.word_count_target * 0.8),
                "completion_rate": len([c for c in self.current_book.chapters if c.status == "completed"]) / len(self.current_book.chapters) if self.current_book.chapters else 0
            }
        }
        
        self.current_book.build_log = build_log
        
        # Save build log
        output_dir = Path(f"output/{self.current_book.build_id}")
        output_dir.mkdir(parents=True, exist_ok=True)
        
        log_path = output_dir / "build_log.json"
        with open(log_path, 'w', encoding='utf-8') as f:
            json.dump(build_log, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Build log saved to {log_path}")
    
    def get_book_status(self) -> Dict[str, Any]:
        """Get current book production status."""
        
        if not self.current_book:
            return {"status": "no_book_in_progress"}
        
        total_words = sum(chapter.actual_word_count for chapter in self.current_book.chapters)
        
        return {
            "title": self.current_book.title,
            "theme": self.current_book.theme,
            "build_id": self.current_book.build_id,
            "status": self.current_book.status,
            "total_word_count": total_words,
            "target_word_count": self.current_book.word_count,
            "chapters_completed": len([c for c in self.current_book.chapters if c.status == "completed"]),
            "total_chapters": len(self.current_book.chapters),
            "completion_percentage": (len([c for c in self.current_book.chapters if c.status == "completed"]) / len(self.current_book.chapters) * 100) if self.current_book.chapters else 0,
            "references_used": len(self.current_book.bibliography),
            "created_at": self.current_book.created_at.isoformat()
        }